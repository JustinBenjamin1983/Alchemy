# report_generator.py - Generate Word document DD report using python-docx
# Fixed to use correct dictionary field mappings from __init__.py

from io import BytesIO
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml


# Severity colors - matching UI status colors
# UI uses: Critical (red), High (orange), Medium (yellow), Positive (green)
COLORS = {
    'Red': RGBColor(185, 28, 28),      # Critical - red-700
    'Amber': RGBColor(234, 88, 12),    # High - orange-600
    'New': RGBColor(234, 88, 12),      # Treat New as High (orange)
    'Green': RGBColor(22, 163, 74),    # Positive - green-600
    'Info': RGBColor(108, 117, 125),   # Gray for informational
    'Medium': RGBColor(202, 138, 4),   # Medium - yellow-600
}

# Display order (most severe first)
SEVERITY_ORDER = ['Red', 'Amber', 'New', 'Green', 'Info']

# Display labels - matching UI terminology
SEVERITY_LABELS = {
    'Red': 'CRITICAL',
    'Amber': 'HIGH',
    'Green': 'POSITIVE',
    'New': 'HIGH',
    'Info': 'INFO',
    'Medium': 'MEDIUM'
}


def generate_dd_report(
    dd_name: str,
    transaction_type: str,
    briefing: str,
    findings: List[Dict],
    documents: List[str],
    synthesis: Dict[str, Any]
) -> BytesIO:
    """Generate a professional Word document DD report.

    Args:
        dd_name: Name of the due diligence project
        transaction_type: Type of transaction (e.g., M&A, Asset Purchase)
        briefing: DD briefing/description
        findings: List of finding dictionaries with keys:
            - id, category, detail, status, phrase, page_number
            - document_name, folder_path, confidence_score
            - finding_type, requires_action, action_priority
            - direct_answer, evidence_quote
        documents: List of document names
        synthesis: Claude synthesis dictionary with:
            - executive_summary, statistics, category_summaries
            - action_items, conclusion
    """

    doc = Document()

    # Set up document styles
    setup_styles(doc)

    # Filter out deleted findings
    active_findings = [f for f in findings if get_status(f) != 'Deleted']

    # DEDUPLICATION: Assign unique report reference IDs to each finding
    # This allows us to reference findings without repeating full content
    active_findings = assign_report_references(active_findings)

    # Add sections
    add_cover_page(doc, dd_name, transaction_type)
    add_table_of_contents(doc)
    add_executive_summary(doc, synthesis, dd_name, transaction_type, active_findings)
    add_deal_assessment(doc, synthesis, active_findings)
    add_financial_exposures(doc, synthesis, active_findings)
    add_deal_blockers(doc, synthesis, active_findings)
    add_conditions_precedent(doc, synthesis, active_findings)
    add_document_status(doc, documents, active_findings)
    add_risk_legend(doc)
    add_risk_sections(doc, active_findings, synthesis)
    add_action_items(doc, active_findings, synthesis)
    add_conclusion(doc, synthesis)

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def get_status(finding: Dict) -> str:
    """Get status string from finding dictionary."""
    status = finding.get('status', 'New')
    # Handle enum values
    if hasattr(status, 'value'):
        return status.value
    return str(status) if status else 'New'


def assign_report_references(findings: List[Dict]) -> List[Dict]:
    """
    Assign unique report reference IDs to findings for deduplication.

    Each finding gets a reference like F001, F002, etc.
    Findings are sorted by severity first so critical issues get lower numbers.
    """
    # Sort by severity (Red first, then Amber, then Green, then Info)
    severity_order = {'Red': 0, 'Amber': 1, 'New': 1, 'Green': 2, 'Info': 3}

    sorted_findings = sorted(
        findings,
        key=lambda f: (
            severity_order.get(get_status(f), 99),
            f.get('category', 'ZZZ'),
            f.get('id', '')
        )
    )

    # Assign reference IDs
    for i, finding in enumerate(sorted_findings, 1):
        finding['_report_ref'] = f"F{i:03d}"

    return sorted_findings


def get_report_ref(finding: Dict) -> str:
    """Get the report reference ID for a finding."""
    return finding.get('_report_ref', 'F???')


def get_finding_summary(finding: Dict, max_length: int = 80) -> str:
    """Get a short summary of a finding for reference sections."""
    title = get_risk_title(finding)
    if len(title) > max_length:
        return title[:max_length - 3] + "..."
    return title


def deduplicate_for_section(findings: List[Dict], section_type: str) -> List[Dict]:
    """
    Return findings appropriate for a given section, avoiding duplication.

    section_type can be:
    - 'executive_summary': Only top 5 critical findings, summary only
    - 'deal_blockers': Only deal blockers, summary only
    - 'conditions_precedent': Only CPs, summary only
    - 'risk_analysis': All findings, full details (main section)
    - 'action_items': Only actionable findings, reference only
    """
    if section_type == 'executive_summary':
        # Top critical findings only
        red_findings = [f for f in findings if get_status(f) == 'Red']
        return red_findings[:5]

    elif section_type == 'deal_blockers':
        return [f for f in findings if f.get('deal_impact') == 'deal_blocker']

    elif section_type == 'conditions_precedent':
        return [f for f in findings if f.get('deal_impact') == 'condition_precedent']

    elif section_type == 'action_items':
        return [
            f for f in findings
            if f.get('requires_action') or get_status(f) == 'Red' or
            f.get('deal_impact') in ('deal_blocker', 'condition_precedent')
        ][:15]

    else:  # 'risk_analysis' - return all
        return findings


def get_category(finding: Dict) -> str:
    """Get category from finding dictionary."""
    return finding.get('category', 'Uncategorized') or 'Uncategorized'


def get_risk_title(finding: Dict) -> str:
    """Get unique finding title - use direct_answer which summarizes the specific finding."""
    # direct_answer is unique per finding and summarizes what was found
    # phrase is the raw text/clause from the document
    # detail is the risk question (shared across multiple findings)
    direct_answer = finding.get('direct_answer', '')
    if direct_answer and len(direct_answer) > 10:
        return direct_answer
    # Fall back to phrase (the actual clause found)
    phrase = finding.get('phrase', '')
    if phrase and len(phrase) > 10:
        return phrase
    # Last resort: use risk detail
    return finding.get('detail', 'Unspecified Risk') or 'Unspecified Risk'


def get_risk_question(finding: Dict) -> str:
    """Get the risk question/category detail (shared across findings of same risk)."""
    return finding.get('detail', '') or ''


def setup_styles(doc: Document):
    """Set up custom styles for the document."""
    styles = doc.styles

    # Heading 1 style
    if 'DD Heading 1' not in [s.name for s in styles]:
        h1_style = styles.add_style('DD Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(12)

    # Heading 2 style
    if 'DD Heading 2' not in [s.name for s in styles]:
        h2_style = styles.add_style('DD Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        h2_style.paragraph_format.space_before = Pt(14)
        h2_style.paragraph_format.space_after = Pt(8)

    # Heading 3 style
    if 'DD Heading 3' not in [s.name for s in styles]:
        h3_style = styles.add_style('DD Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)
        h3_style.paragraph_format.space_before = Pt(10)
        h3_style.paragraph_format.space_after = Pt(6)


def add_cover_page(doc: Document, dd_name: str, transaction_type: str):
    """Add a professional cover page."""
    # Add some spacing at top
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PRELIMINARY DUE DILIGENCE REPORT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph()

    # Project name
    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project.add_run(dd_name)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

    doc.add_paragraph()

    # Transaction type
    trans = doc.add_paragraph()
    trans.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = trans.add_run(f"Transaction Type: {transaction_type}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

    # Add spacing
    for _ in range(8):
        doc.add_paragraph()

    # Date and confidentiality
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("CONFIDENTIAL - ATTORNEY-CLIENT PRIVILEGED")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

    # Page break
    doc.add_page_break()


def add_table_of_contents(doc: Document):
    """Add a table of contents placeholder."""
    doc.add_paragraph("TABLE OF CONTENTS", style='DD Heading 1')

    toc_items = [
        ("1. Executive Summary", 3),
        ("2. Deal Assessment", 4),
        ("3. Financial Exposures", 5),
        ("4. Deal Blockers", 6),
        ("5. Conditions Precedent", 7),
        ("6. Documents Reviewed", 8),
        ("7. Risk Classification Legend", 9),
        ("8. Risk Analysis by Category", 10),
        ("9. Action Items & Queries", 0),
        ("10. Conclusion & Recommendations", 0),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        if page > 0:
            p.add_run(f"{'.' * 60}{page}").font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)

    doc.add_page_break()


def add_executive_summary(doc: Document, synthesis: Dict, dd_name: str, transaction_type: str, findings: List[Dict]):
    """Add the executive summary section."""
    doc.add_paragraph("1. EXECUTIVE SUMMARY", style='DD Heading 1')

    # Overview paragraph from synthesis
    overview = synthesis.get('executive_summary', {}).get('overview', '')
    if overview:
        doc.add_paragraph(overview)
    else:
        doc.add_paragraph(
            f"This preliminary due diligence report presents findings from the review of {dd_name}. "
            f"A total of {len(findings)} findings were identified across the documents analyzed."
        )

    # Key statistics table
    doc.add_paragraph("Key Metrics", style='DD Heading 2')

    # Calculate stats from findings
    red_count = sum(1 for f in findings if get_status(f) == 'Red')
    amber_count = sum(1 for f in findings if get_status(f) in ['Amber', 'New'])
    green_count = sum(1 for f in findings if get_status(f) == 'Green')
    action_count = sum(1 for f in findings if f.get('requires_action'))

    # Enhanced: Get deal blockers and conditions precedent counts from synthesis (if available)
    # This ensures counts match what Pass 4 identified, not raw finding counts
    synthesis_deal_blockers = synthesis.get('deal_blockers', {})
    synthesis_cps = synthesis.get('conditions_precedent', {})

    # Use synthesis counts if available (from Pass 4 analysis)
    if isinstance(synthesis_deal_blockers, dict) and 'blockers' in synthesis_deal_blockers:
        deal_blocker_count = len(synthesis_deal_blockers.get('blockers', []))
    else:
        deal_blocker_count = len(synthesis_deal_blockers) if isinstance(synthesis_deal_blockers, list) else 0

    if isinstance(synthesis_cps, dict) and 'conditions' in synthesis_cps:
        cp_count = len(synthesis_cps.get('conditions', []))
    else:
        cp_count = len(synthesis_cps) if isinstance(synthesis_cps, list) else 0

    cross_doc_count = sum(1 for f in findings if f.get('is_cross_doc'))

    # Calculate total financial exposure
    total_exposure = 0
    exposure_currency = 'ZAR'
    for f in findings:
        fin_exp = f.get('financial_exposure')
        if fin_exp and fin_exp.get('amount'):
            total_exposure += fin_exp['amount']
            if fin_exp.get('currency'):
                exposure_currency = fin_exp['currency']

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row - matching UI terminology
    hdr_cells = table.rows[0].cells
    headers = ['Total Findings', 'Critical', 'High', 'Action Required']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], "1a365d")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data row
    row = table.add_row()
    row.cells[0].text = str(len(findings))
    row.cells[1].text = str(red_count)
    row.cells[2].text = str(amber_count)
    row.cells[3].text = str(action_count)

    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Color code critical findings cell
    if red_count > 0:
        set_cell_shading(row.cells[1], "f8d7da")

    doc.add_paragraph()

    # Enhanced: Deal Impact Summary (only if we have deal impact data)
    if deal_blocker_count > 0 or cp_count > 0 or total_exposure > 0:
        doc.add_paragraph("Deal Impact Summary", style='DD Heading 2')

        impact_table = doc.add_table(rows=1, cols=4)
        impact_table.style = 'Table Grid'
        impact_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = impact_table.rows[0].cells
        headers = ['Deal Blockers', 'Conditions Precedent', 'Cross-Doc Findings', 'Total Financial Exposure']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(hdr_cells[i], "1a365d")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Data row
        row = impact_table.add_row()
        row.cells[0].text = str(deal_blocker_count)
        row.cells[1].text = str(cp_count)
        row.cells[2].text = str(cross_doc_count)
        row.cells[3].text = f"{exposure_currency} {total_exposure:,.0f}" if total_exposure > 0 else "N/A"

        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Color code deal blockers
        if deal_blocker_count > 0:
            set_cell_shading(row.cells[0], "f8d7da")  # Red background for deal blockers
        if total_exposure > 0:
            set_cell_shading(row.cells[3], "fff3cd")  # Amber background for financial exposure

        doc.add_paragraph()

    # Key findings summary from synthesis
    doc.add_paragraph("Key Findings", style='DD Heading 2')

    key_findings = synthesis.get('executive_summary', {}).get('key_findings', [])
    if key_findings:
        for finding in key_findings[:5]:  # Top 5 key findings
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(finding)
    else:
        # Generate from actual Red findings - use REFERENCES to avoid duplication
        red_findings = deduplicate_for_section(findings, 'executive_summary')
        if red_findings:
            doc.add_paragraph(
                "The following critical findings require immediate attention. "
                "See Section 8 (Risk Analysis) for full details."
            )
            for f in red_findings:
                p = doc.add_paragraph(style='List Bullet')
                # Include reference ID for cross-referencing
                ref = get_report_ref(f)
                run = p.add_run(f"[{ref}] ")
                run.bold = True
                run.font.color.rgb = COLORS.get(get_status(f), COLORS['Amber'])
                p.add_run(f"{get_category(f)}: {get_finding_summary(f, 100)}")
        else:
            doc.add_paragraph("No critical findings identified.")

    doc.add_page_break()


def add_deal_assessment(doc: Document, synthesis: Dict, findings: List[Dict]):
    """Add the deal assessment section."""
    doc.add_paragraph("2. DEAL ASSESSMENT", style='DD Heading 1')

    deal_assessment = synthesis.get('deal_assessment', {})

    # Overall Viability
    doc.add_paragraph("Overall Deal Viability", style='DD Heading 2')

    viability = deal_assessment.get('overall_viability', 'Assessment pending further analysis')
    p = doc.add_paragraph()
    run = p.add_run("Recommendation: ")
    run.bold = True

    # Color code based on recommendation
    viability_text = str(viability)
    viability_run = p.add_run(viability_text)
    viability_run.bold = True
    if 'No Go' in viability_text:
        viability_run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)  # Red
    elif 'Conditional' in viability_text or 'Caution' in viability_text:
        viability_run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)  # Orange
    else:
        viability_run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)  # Green

    # Key Concerns
    doc.add_paragraph("Key Concerns", style='DD Heading 2')
    key_concerns = deal_assessment.get('key_concerns', [])
    if key_concerns:
        for concern in key_concerns[:7]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(str(concern)[:200])
    else:
        doc.add_paragraph("No major concerns identified.")

    # Mitigating Factors
    doc.add_paragraph("Mitigating Factors", style='DD Heading 2')
    mitigating = deal_assessment.get('mitigating_factors', [])
    if mitigating:
        for factor in mitigating[:5]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(str(factor)[:200])
    else:
        doc.add_paragraph("Further review may identify mitigating factors.")

    # Negotiation Points
    doc.add_paragraph("Key Negotiation Points", style='DD Heading 2')
    negotiation = deal_assessment.get('negotiation_points', [])
    if negotiation:
        for point in negotiation[:7]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(str(point)[:200])
    else:
        doc.add_paragraph("Standard negotiation points apply.")

    # Deal Structure Recommendations
    doc.add_paragraph("Deal Structure Recommendations", style='DD Heading 2')
    structure = deal_assessment.get('deal_structure_recommendations', '')
    if structure:
        doc.add_paragraph(str(structure))
    else:
        doc.add_paragraph("Consider standard deal protections including warranties, indemnities, and escrow arrangements.")

    doc.add_page_break()


def add_financial_exposures(doc: Document, synthesis: Dict, findings: List[Dict]):
    """Add the financial exposures section."""
    doc.add_paragraph("3. FINANCIAL EXPOSURES", style='DD Heading 1')

    fin_exposures = synthesis.get('financial_exposures', {})

    # Total Exposure
    doc.add_paragraph("Total Quantified Exposure", style='DD Heading 2')
    total = fin_exposures.get('total_quantified_exposure', 'Not quantified')

    p = doc.add_paragraph()
    p.add_run("Total: ").bold = True
    run = p.add_run(str(total))
    run.bold = True
    if total and 'Not' not in str(total):
        run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)  # Red for monetary amounts
        run.font.size = Pt(14)

    # Exposure Breakdown Table
    breakdown = fin_exposures.get('exposure_breakdown', [])
    if breakdown:
        doc.add_paragraph("Exposure Breakdown", style='DD Heading 2')

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        headers = ['Category', 'Amount', 'Description', 'Source', 'Mitigation']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(hdr_cells[i], "1a365d")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for exp in breakdown[:15]:
            row = table.add_row()
            row.cells[0].text = str(exp.get('category', ''))[:20]
            row.cells[1].text = str(exp.get('amount', ''))[:20]
            row.cells[2].text = str(exp.get('description', ''))[:50]
            row.cells[3].text = str(exp.get('source_document', ''))[:25]
            row.cells[4].text = str(exp.get('mitigation', ''))[:40]

            # Color code amount cell
            set_cell_shading(row.cells[1], "fff3cd")

    # Unquantified Risks
    doc.add_paragraph()
    doc.add_paragraph("Unquantified Risks", style='DD Heading 2')
    unquantified = fin_exposures.get('unquantified_risks', [])
    if unquantified:
        for risk in unquantified[:10]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(str(risk)[:200])
    else:
        doc.add_paragraph("All identified financial risks have been quantified where possible.")

    # Recommended Provisions
    doc.add_paragraph("Recommended Financial Provisions", style='DD Heading 2')
    provisions = fin_exposures.get('recommended_provisions', '')
    if provisions:
        doc.add_paragraph(str(provisions))
    else:
        doc.add_paragraph("Standard indemnity and escrow provisions recommended based on transaction value.")

    doc.add_page_break()


def add_deal_blockers(doc: Document, synthesis: Dict, findings: List[Dict]):
    """Add the deal blockers section with deduplication via references."""
    doc.add_paragraph("4. DEAL BLOCKERS", style='DD Heading 1')

    blockers_data = synthesis.get('deal_blockers', {})

    # Summary
    summary = blockers_data.get('summary', '')
    if summary:
        p = doc.add_paragraph()
        run = p.add_run(str(summary))
        run.italic = True

    blockers = blockers_data.get('blockers', [])

    # Get deal blocker findings using deduplication helper
    blocker_findings = deduplicate_for_section(findings, 'deal_blockers')

    if not blockers and blocker_findings:
        # Generate blockers from findings with references
        for f in blocker_findings[:10]:
            blockers.append({
                'report_ref': get_report_ref(f),
                'title': get_finding_summary(f, 120),
                'category': f.get('category', 'General'),
                'source_document': f.get('document_name', 'Unknown'),
                'resolution_path': f.get('resolution_path', 'Requires resolution'),
                'criticality': 'Deal-blocking issue - must be resolved before closing',
                # Store original finding for reference
                '_finding': f
            })

    if blockers:
        doc.add_paragraph()
        doc.add_paragraph("Identified Deal Blockers", style='DD Heading 2')

        # Add cross-reference note
        ref_note = doc.add_paragraph()
        ref_note.add_run(
            "Note: Reference IDs (e.g., F001) link to full details in Section 8 (Risk Analysis). "
        ).font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

        for i, blocker in enumerate(blockers[:10], 1):
            # Blocker header with report reference
            p = doc.add_paragraph()

            # Add report reference if available
            report_ref = blocker.get('report_ref', '')
            if report_ref:
                run = p.add_run(f"[{report_ref}] ")
                run.bold = True
                run.font.color.rgb = RGBColor(0x27, 0x63, 0x89)

            run = p.add_run(f"Blocker {i}: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

            title = blocker.get('title', 'Unspecified blocker')
            p.add_run(str(title)[:150]).bold = True

            # Category and source (compact reference line)
            ref = doc.add_paragraph()
            cat = blocker.get('category', '')
            source = blocker.get('source_document', '')
            ref.add_run(f"Category: {cat} | Source: {source}").font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

            # Resolution path (brief - full details in Section 8)
            resolution = blocker.get('resolution_path', '')
            if resolution and resolution != 'Requires resolution':
                res_p = doc.add_paragraph()
                res_p.add_run("Resolution Path: ").bold = True
                res_p.add_run(str(resolution)[:150])

            # Financial exposure if available (from original finding)
            original_finding = blocker.get('_finding')
            if original_finding:
                financial_exposure = original_finding.get('financial_exposure')
                if financial_exposure and financial_exposure.get('amount'):
                    fin_p = doc.add_paragraph()
                    fin_p.add_run("Financial Exposure: ").bold = True
                    amount = financial_exposure['amount']
                    currency = financial_exposure.get('currency', 'ZAR')
                    run = fin_p.add_run(f"{currency} {amount:,.0f}")
                    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
                    run.bold = True

            doc.add_paragraph()

        # Resolution Timeline
        timeline = blockers_data.get('resolution_timeline', '')
        if timeline:
            doc.add_paragraph("Estimated Resolution Timeline", style='DD Heading 2')
            doc.add_paragraph(str(timeline))
    else:
        p = doc.add_paragraph()
        run = p.add_run("No deal blockers identified.")
        run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)
        run.bold = True

        doc.add_paragraph()
        doc.add_paragraph(
            "While no absolute deal blockers were identified, please review the Risk Analysis section "
            "for items that may require attention before transaction completion."
        )

    doc.add_page_break()


def add_conditions_precedent(doc: Document, synthesis: Dict, findings: List[Dict]):
    """Add the conditions precedent section with deduplication via references."""
    doc.add_paragraph("5. CONDITIONS PRECEDENT", style='DD Heading 1')

    cp_data = synthesis.get('conditions_precedent', {})

    # Summary
    summary = cp_data.get('summary', '')
    if summary:
        p = doc.add_paragraph()
        run = p.add_run(str(summary))
        run.italic = True

    conditions = cp_data.get('conditions', [])

    # Get CP findings using deduplication helper
    cp_findings = deduplicate_for_section(findings, 'conditions_precedent')

    if not conditions and cp_findings:
        # Generate conditions from findings with references
        for f in cp_findings[:10]:
            conditions.append({
                'report_ref': get_report_ref(f),
                'title': get_finding_summary(f, 80),
                'category': f.get('category', 'General'),
                'responsible_party': f.get('responsible_party', 'TBD'),
                'estimated_timeline': 'Before closing',
                'risk_if_not_met': 'Transaction cannot close'
            })

    if conditions:
        # Create conditions table
        doc.add_paragraph()
        doc.add_paragraph("Conditions to be Satisfied", style='DD Heading 2')

        # Add cross-reference note
        ref_note = doc.add_paragraph()
        ref_note.add_run(
            "Note: Reference IDs link to full details in Section 8 (Risk Analysis). "
        ).font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        headers = ['Ref', 'Condition', 'Category', 'Responsible', 'Timeline']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(hdr_cells[i], "1a365d")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for cp in conditions[:15]:
            row = table.add_row()
            # Add report reference in first column
            report_ref = cp.get('report_ref', '-')
            row.cells[0].text = report_ref
            row.cells[1].text = str(cp.get('title', ''))[:50]
            row.cells[2].text = str(cp.get('category', ''))[:15]
            row.cells[3].text = str(cp.get('responsible_party', ''))[:12]
            row.cells[4].text = str(cp.get('estimated_timeline', ''))[:15]

            # Color code reference column
            if report_ref and report_ref != '-':
                set_cell_shading(row.cells[0], "e3f2fd")  # Light blue for references

        # Critical Path Items
        doc.add_paragraph()
        critical_path = cp_data.get('critical_path_items', [])
        if critical_path:
            doc.add_paragraph("Critical Path Items", style='DD Heading 2')
            for item in critical_path[:7]:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(str(item)[:200])
                run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
    else:
        p = doc.add_paragraph()
        p.add_run("Standard conditions precedent will apply to this transaction, including:")

        standard_cps = [
            "Completion of satisfactory due diligence",
            "Execution of definitive agreements",
            "Receipt of necessary regulatory approvals",
            "No material adverse change",
            "Delivery of closing documentation"
        ]
        for cp in standard_cps:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(cp)

    doc.add_page_break()


def add_document_status(doc: Document, documents: List[str], findings: List[Dict]):
    """Add document review status section."""
    doc.add_paragraph("6. DOCUMENTS REVIEWED", style='DD Heading 1')

    doc.add_paragraph(
        f"A total of {len(documents)} documents were reviewed as part of this due diligence exercise. "
        "The following table summarizes the documents analyzed and their risk profiles."
    )

    # Create document summary table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header
    headers = ['Document Name', 'Findings', 'Highest Risk', 'Status']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], "1a365d")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Group findings by document
    doc_findings = {}
    for f in findings:
        doc_name = f.get('document_name', 'Unknown')
        if doc_name not in doc_findings:
            doc_findings[doc_name] = []
        doc_findings[doc_name].append(f)

    # Add rows for each document
    for doc_name in documents[:20]:  # Limit to first 20 docs
        row = table.add_row()
        row.cells[0].text = doc_name[:50] + ('...' if len(doc_name) > 50 else '')

        doc_f = doc_findings.get(doc_name, [])
        row.cells[1].text = str(len(doc_f))

        # Determine highest risk using status values
        statuses = [get_status(f) for f in doc_f]
        if 'Red' in statuses:
            highest = 'Red'
            color = "f8d7da"
        elif 'Amber' in statuses or 'New' in statuses:
            highest = 'Amber'
            color = "fff3cd"
        elif doc_f:
            highest = 'Green'
            color = "d4edda"
        else:
            highest = 'N/A'
            color = "e2e3e5"

        row.cells[2].text = highest
        set_cell_shading(row.cells[2], color)

        row.cells[3].text = 'Reviewed' if doc_f else 'No findings'

    if len(documents) > 20:
        doc.add_paragraph(f"... and {len(documents) - 20} additional documents")

    doc.add_page_break()


def add_risk_legend(doc: Document):
    """Add risk classification legend matching UI colors."""
    doc.add_paragraph("7. RISK CLASSIFICATION LEGEND", style='DD Heading 1')

    doc.add_paragraph(
        "Findings are classified according to the following risk levels based on their "
        "potential impact on the transaction:"
    )

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # Matching UI: Critical (red), High (orange/peach), Medium (yellow), Positive (green)
    legend_data = [
        ('Severity', 'Meaning'),
        ('CRITICAL', 'Critical risk requiring immediate attention. May be deal-breaker or require significant negotiation/price adjustment.'),
        ('HIGH', 'Material risk requiring further investigation or disclosure. Should be addressed in SPA warranties or indemnities.'),
        ('MEDIUM', 'Moderate risk that should be noted. May require standard contract protections.'),
        ('POSITIVE', 'Low risk, positive finding, or standard commercial term. Note for awareness and completeness only.'),
    ]

    # Colors matching UI: Critical=red, High=orange/peach, Medium=yellow, Positive=green
    row_colors = ['1a365d', 'fecaca', 'fed7aa', 'fef08a', 'bbf7d0']

    for i, (severity, meaning) in enumerate(legend_data):
        row = table.rows[i]
        row.cells[0].text = severity
        row.cells[1].text = meaning

        if i == 0:  # Header row
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, row_colors[0])
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_shading(row.cells[0], row_colors[i])

    doc.add_paragraph()
    doc.add_page_break()


def add_risk_sections(doc: Document, findings: List[Dict], synthesis: Dict):
    """Add detailed risk sections by category."""
    doc.add_paragraph("8. RISK ANALYSIS BY CATEGORY", style='DD Heading 1')

    # Group findings by category
    categories = {}
    for f in findings:
        cat = get_category(f)
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(f)

    # If no findings at all
    if not categories:
        doc.add_paragraph("No material risks identified in the documents reviewed.")
        doc.add_page_break()
        return

    # Get category summaries from synthesis
    category_summaries = synthesis.get('category_summaries', {})

    # Sort categories by severity (most Red findings first)
    def category_severity(cat_name):
        cat_findings = categories.get(cat_name, [])
        red_count = sum(1 for f in cat_findings if get_status(f) == 'Red')
        amber_count = sum(1 for f in cat_findings if get_status(f) in ['Amber', 'New'])
        return (-red_count, -amber_count, cat_name)

    sorted_categories = sorted(categories.keys(), key=category_severity)

    section_num = 1
    for cat_name in sorted_categories:
        cat_findings = categories[cat_name]

        # Category heading
        doc.add_paragraph(f"8.{section_num} {cat_name}", style='DD Heading 2')

        # Category summary from synthesis
        summary = category_summaries.get(cat_name, {}).get('summary', '')
        if summary:
            p = doc.add_paragraph()
            p.add_run(summary).italic = True

        # Count by severity using status values
        red_count = sum(1 for f in cat_findings if get_status(f) == 'Red')
        amber_count = sum(1 for f in cat_findings if get_status(f) in ['Amber', 'New'])
        green_count = sum(1 for f in cat_findings if get_status(f) == 'Green')
        info_count = sum(1 for f in cat_findings if get_status(f) == 'Info')

        p = doc.add_paragraph()
        p.add_run(f"Total Findings: {len(cat_findings)} ")
        p.add_run("(")
        parts = []
        if red_count:
            run = p.add_run(f"{red_count} Red")
            run.font.color.rgb = COLORS['Red']
            run.bold = True
            parts.append(True)
        if amber_count:
            if parts:
                p.add_run(", ")
            run = p.add_run(f"{amber_count} Amber")
            run.font.color.rgb = COLORS['Amber']
            parts.append(True)
        if green_count:
            if parts:
                p.add_run(", ")
            run = p.add_run(f"{green_count} Green")
            run.font.color.rgb = COLORS['Green']
            parts.append(True)
        if info_count:
            if parts:
                p.add_run(", ")
            run = p.add_run(f"{info_count} Info")
            run.font.color.rgb = COLORS['Info']
        p.add_run(")")

        # Sort findings by severity within category
        sorted_findings = sorted(
            cat_findings,
            key=lambda x: SEVERITY_ORDER.index(get_status(x)) if get_status(x) in SEVERITY_ORDER else 99
        )

        # Add each finding
        for finding in sorted_findings[:15]:  # Limit to 15 per category
            add_risk_finding(doc, finding)

        if len(cat_findings) > 15:
            doc.add_paragraph(f"... and {len(cat_findings) - 15} additional findings in this category")

        doc.add_paragraph()
        section_num += 1

    doc.add_page_break()


def add_risk_finding(doc: Document, finding: Dict):
    """Add a single risk finding to the document using correct field mappings."""
    status = get_status(finding)
    label = SEVERITY_LABELS.get(status, 'AMBER')
    color = COLORS.get(status, COLORS['Amber'])

    # Finding header with report reference, severity badge, and finding title
    p = doc.add_paragraph()

    # Report reference ID (for cross-referencing from other sections)
    report_ref = get_report_ref(finding)
    if report_ref and report_ref != 'F???':
        run = p.add_run(f"[{report_ref}] ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x27, 0x63, 0x89)  # Teal for reference

    # Severity indicator
    run = p.add_run(f"[{label}] ")
    run.bold = True
    run.font.color.rgb = color

    # Deal impact badge (if available)
    deal_impact = finding.get('deal_impact')
    if deal_impact and deal_impact not in ('none', 'noted'):
        impact_labels = {
            'deal_blocker': 'DEAL BLOCKER',
            'condition_precedent': 'CP',
            'price_chip': 'PRICE CHIP',
            'warranty_indemnity': 'W&I',
            'post_closing': 'POST-CLOSING'
        }
        impact_label = impact_labels.get(deal_impact, deal_impact.upper())
        run = p.add_run(f"[{impact_label}] ")
        run.bold = True
        if deal_impact == 'deal_blocker':
            run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)  # Red for deal blockers
        elif deal_impact == 'condition_precedent':
            run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)  # Orange for CPs
        else:
            run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

    # Finding title - now uses direct_answer which is unique per finding
    finding_title = get_risk_title(finding)
    title_run = p.add_run(finding_title[:200])
    title_run.bold = True

    # Show the risk question as context (only if different from title)
    risk_question = get_risk_question(finding)
    if risk_question and risk_question != finding_title and len(risk_question) > 10:
        context_para = doc.add_paragraph()
        context_para.add_run("Risk Area: ").bold = True
        context_para.add_run(risk_question[:150])

    # Show the specific clause/phrase from the document (only if different from title)
    phrase = finding.get('phrase', '')
    if phrase and phrase != finding_title and len(phrase) > 10:
        clause_para = doc.add_paragraph()
        clause_para.add_run("Clause: ").bold = True
        clause_para.add_run(phrase[:400] + ('...' if len(phrase) > 400 else ''))

    # Clause reference (if available and different from page number)
    clause_ref = finding.get('clause_reference', '')
    if clause_ref and clause_ref != finding.get('page_number', ''):
        ref_para = doc.add_paragraph()
        ref_para.add_run("Reference: ").bold = True
        ref_para.add_run(clause_ref)

    # Document reference - handle cross-doc findings
    is_cross_doc = finding.get('is_cross_doc', False)
    doc_name = finding.get('document_name', 'Unknown')
    cross_doc_source = finding.get('cross_doc_source', '')
    page = finding.get('page_number', 'N/A')

    ref = doc.add_paragraph()
    if is_cross_doc and cross_doc_source:
        run = ref.add_run(f"Cross-Document Analysis: {cross_doc_source}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x27, 0x63, 0x89)  # Teal for cross-doc
        run.italic = True
        run.bold = True
    else:
        run = ref.add_run(f"Source: {doc_name} | Page: {page}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
        run.italic = True

    # Financial exposure (if available) - enhanced for calculation engine results
    financial_exposure = finding.get('financial_exposure')
    calculated_exposure = finding.get('calculated_exposure')

    if financial_exposure and financial_exposure.get('amount'):
        amount = financial_exposure['amount']
        currency = financial_exposure.get('currency', 'ZAR')
        formula_id = financial_exposure.get('formula_id', '')
        confidence = financial_exposure.get('confidence', 0)
        validation_passed = financial_exposure.get('validation_passed', True)
        warnings = financial_exposure.get('warnings', [])

        # Main exposure line
        fin_para = doc.add_paragraph()
        fin_para.add_run("Financial Exposure: ").bold = True
        run = fin_para.add_run(f"{currency} {amount:,.0f}")
        run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)  # Red
        run.bold = True

        # Add confidence level if from calculation engine
        if confidence > 0:
            conf_level = "HIGH" if confidence >= 0.9 else "MEDIUM" if confidence >= 0.7 else "LOW"
            run = fin_para.add_run(f"  (Confidence: {conf_level})")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

        # Add warning indicator if validation issues
        if not validation_passed or warnings:
            run = fin_para.add_run("  ⚠️")
            run.font.size = Pt(9)

        # Show formula and calculation method if available
        if formula_id:
            formula_para = doc.add_paragraph()
            formula_para.paragraph_format.left_indent = Inches(0.3)
            run = formula_para.add_run(f"Formula: {formula_id}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x27, 0x63, 0x89)  # Teal
            run.italic = True

        # Show calculation working
        calculation = financial_exposure.get('calculation', '')
        if calculation:
            calc_para = doc.add_paragraph()
            calc_para.paragraph_format.left_indent = Inches(0.3)
            run = calc_para.add_run(f"Calculation: {calculation[:200]}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
            run.italic = True

        # Show warnings from calculation engine if any
        if warnings:
            for warning in warnings[:2]:
                warn_para = doc.add_paragraph()
                warn_para.paragraph_format.left_indent = Inches(0.3)
                run = warn_para.add_run(f"⚠️ {warning[:100]}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)  # Orange

    # Evidence quote if available and different from phrase
    evidence = finding.get('evidence_quote', '')
    if evidence and evidence != phrase and len(evidence) > 20:
        quote_para = doc.add_paragraph()
        quote_para.paragraph_format.left_indent = Inches(0.5)
        run = quote_para.add_run(f'"{evidence[:300]}{"..." if len(evidence) > 300 else ""}"')
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(0x5d, 0x6d, 0x7e)

    # Confidence score if available
    confidence = finding.get('confidence_score')
    if confidence is not None:
        conf_para = doc.add_paragraph()
        conf_para.add_run("Confidence: ").bold = True
        conf_para.add_run(f"{int(float(confidence) * 100)}%")

    # Action required indicator
    if finding.get('requires_action'):
        action_p = doc.add_paragraph()
        run = action_p.add_run("! Action Required")
        run.bold = True
        run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        run.font.size = Pt(9)

    # Add some spacing
    doc.add_paragraph()


def add_action_items(doc: Document, findings: List[Dict], synthesis: Dict):
    """Add priority action items section with deduplication via references."""
    doc.add_paragraph("9. ACTION ITEMS & QUERIES", style='DD Heading 1')

    doc.add_paragraph(
        "The following items require attention and should be addressed as part of the transaction process. "
        "Reference IDs link to full details in Section 8 (Risk Analysis)."
    )

    # Get action items from synthesis or generate from findings
    action_items = synthesis.get('action_items', [])

    if not action_items:
        # Use deduplication helper to get action-worthy findings
        action_findings = deduplicate_for_section(findings, 'action_items')

        # Sort by deal impact priority (deal blockers first, then CPs, then red status)
        def action_priority_sort(f):
            if f.get('deal_impact') == 'deal_blocker':
                return 0
            elif f.get('deal_impact') == 'condition_precedent':
                return 1
            elif get_status(f) == 'Red':
                return 2
            return 3

        action_findings.sort(key=action_priority_sort)

        for f in action_findings[:15]:
            deal_impact = f.get('deal_impact', '')
            if deal_impact == 'deal_blocker':
                priority = 'critical'
            elif deal_impact == 'condition_precedent':
                priority = 'high'
            elif get_status(f) == 'Red':
                priority = 'high'
            else:
                priority = f.get('action_priority', 'medium')

            action_items.append({
                'report_ref': get_report_ref(f),
                'priority': priority,
                'description': get_finding_summary(f, 60),
                'category': get_category(f),
                'source': f.get('document_name', 'Unknown')[:30],
                'deal_impact': deal_impact
            })

    if action_items:
        # Create action items table with reference column
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        headers = ['Ref', 'Priority', 'Impact', 'Action Item', 'Category', 'Source']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(hdr_cells[i], "1a365d")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for item in action_items[:15]:
            row = table.add_row()
            priority = str(item.get('priority', 'medium')).upper()
            deal_impact = item.get('deal_impact', '')
            report_ref = item.get('report_ref', '-')

            # Map deal impact to display label
            impact_labels = {
                'deal_blocker': 'BLOCKER',
                'condition_precedent': 'CP',
                'price_chip': 'PRICE',
                'warranty_indemnity': 'W&I',
                'post_closing': 'POST'
            }
            impact_display = impact_labels.get(deal_impact, '-')

            row.cells[0].text = report_ref
            row.cells[1].text = priority
            row.cells[2].text = impact_display
            row.cells[3].text = item.get('description', item.get('query', ''))[:55]
            row.cells[4].text = item.get('category', '')[:12]
            row.cells[5].text = item.get('source', item.get('document', ''))[:20]

            # Color code reference column
            if report_ref and report_ref != '-':
                set_cell_shading(row.cells[0], "e3f2fd")  # Light blue for references

            # Color code priority
            if priority == 'CRITICAL':
                set_cell_shading(row.cells[1], "c00020")  # Dark red
                if row.cells[1].paragraphs[0].runs:
                    row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            elif priority == 'HIGH':
                set_cell_shading(row.cells[1], "f8d7da")
            elif priority in ['MEDIUM', 'AMBER']:
                set_cell_shading(row.cells[1], "fff3cd")

            # Color code deal impact
            if deal_impact == 'deal_blocker':
                set_cell_shading(row.cells[2], "f8d7da")
            elif deal_impact == 'condition_precedent':
                set_cell_shading(row.cells[2], "fff3cd")
    else:
        doc.add_paragraph("Review individual findings above for specific recommendations.")

    doc.add_page_break()


def add_conclusion(doc: Document, synthesis: Dict):
    """Add conclusion and recommendations section."""
    doc.add_paragraph("10. CONCLUSION & RECOMMENDATIONS", style='DD Heading 1')

    # Overall assessment
    doc.add_paragraph("Overall Assessment", style='DD Heading 2')

    conclusion = synthesis.get('conclusion', {})
    assessment = conclusion.get('overall_assessment', '')
    if assessment:
        doc.add_paragraph(assessment)
    else:
        doc.add_paragraph(
            "Based on our review of the available documentation, we have identified the key risks and issues "
            "outlined in this report. Further investigation may be required in certain areas."
        )

    # Recommendations
    doc.add_paragraph("Key Recommendations", style='DD Heading 2')

    recommendations = conclusion.get('recommendations', [])
    if recommendations:
        for i, rec in enumerate(recommendations, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. ").bold = True
            p.add_run(rec)
    else:
        doc.add_paragraph("1. Review all Red findings with legal counsel", style='List Number')
        doc.add_paragraph("2. Address Amber items in transaction documentation", style='List Number')
        doc.add_paragraph("3. Consider further investigation of flagged areas", style='List Number')

    # Next steps
    doc.add_paragraph("Next Steps", style='DD Heading 2')

    next_steps = conclusion.get('next_steps', [])
    if next_steps:
        for step in next_steps:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(step)
    else:
        doc.add_paragraph("Complete review of outstanding documents", style='List Bullet')
        doc.add_paragraph("Discuss findings with transaction team", style='List Bullet')
        doc.add_paragraph("Prepare queries for counterparty", style='List Bullet')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("-- End of Report --")
    run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    run.italic = True

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("Disclaimer: ").bold = True
    disclaimer.add_run(
        "This preliminary report is based on automated analysis of provided documents and should be "
        "reviewed by qualified legal professionals before reliance. It does not constitute legal advice. "
        "Additional documents and information may materially affect the risk assessment."
    )

    # Generated by note
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen.add_run(f"Report generated by Alchemy Legal AI on {datetime.now().strftime('%d %B %Y at %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)


def set_cell_shading(cell, color_hex: str):
    """Set the background shading of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)
