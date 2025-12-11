# DDQuestionsExport/__init__.py
import logging
import os
import json
import io
from datetime import datetime
import azure.functions as func
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from shared.utils import auth_get_email
from shared.session import transactional_session
from shared.models import DDQuestion, DDQuestionReferencedDoc, DueDiligence
from sqlalchemy.orm import joinedload
import re

def clean_markdown_for_word(text):
    """Convert markdown text to plain text suitable for Word document"""
    if not text:
        return ""
    
    # Remove markdown headers
    text = re.sub(r'^#+\s*(.*)$', r'\1', text, flags=re.MULTILINE)
    
    # Convert bold markdown to plain text (we'll handle formatting separately)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Convert italic markdown to plain text
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # Convert code blocks to plain text
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # Remove bullet points markdown
    text = re.sub(r'^\s*[\*\-\+]\s+', '• ', text, flags=re.MULTILINE)
    
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = text.strip()
    
    return text

def extract_confidence_level(text):
    """Extract confidence indicators from text"""
    confidence_pattern = r'(High Confidence \(\d+\-\d+%\)|Medium Confidence \(\d+\-\d+%\)|Low Confidence \(\d+\-\d+%\)|Uncertain \(<\d+%\))'
    matches = re.findall(confidence_pattern, text)
    return matches

def add_header_footer(doc, dd_title):
    """Add header and footer to the document"""
    # Add header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"Due Diligence Q&A Report - {dd_title}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add footer with page numbers and date
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_word_document(questions_data, dd_title="Due Diligence"):
    """Create a professional Word document from Q&A data"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header and footer
    add_header_footer(doc, dd_title)
    
    # Title
    title = doc.add_heading(f'{dd_title}\nQuestions & Answers Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary statistics
    total_questions = len(questions_data)
    answered_questions = len([q for q in questions_data if q.get('answer')])
    unanswered_questions = total_questions - answered_questions
    
    # Add summary paragraph
    summary_para = doc.add_paragraph()
    summary_para.add_run('Report Summary: ').bold = True
    summary_para.add_run(f'This report contains {total_questions} questions, of which {answered_questions} have been answered and {unanswered_questions} could not be answered based on the available documentation.')
    
    # Add generation date
    date_para = doc.add_paragraph()
    date_para.add_run('Generated: ').bold = True
    date_para.add_run(datetime.now().strftime('%B %d, %Y at %I:%M %p'))
    
    # Add page break
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', level=1)
    
    for i, question in enumerate(questions_data, 1):
        toc_para = doc.add_paragraph()
        toc_para.add_run(f'{i}. ').bold = True
        question_preview = question['question'][:80] + ('...' if len(question['question']) > 80 else '')
        toc_para.add_run(question_preview)
        
        # Add context info
        if question.get('document_name'):
            toc_para.add_run(f' (Document: {question["document_name"]})')
        elif question.get('folder_name'):
            toc_para.add_run(f' (Folder: {question["folder_name"]})')
    
    # Add page break before Q&A section
    doc.add_page_break()
    
    # Questions and Answers section
    qa_heading = doc.add_heading('Questions & Answers', level=1)
    
    for i, question in enumerate(questions_data, 1):
        # Question number and text
        q_heading = doc.add_heading(f'Question {i}', level=2)
        
        # Add question text
        q_para = doc.add_paragraph()
        q_para.add_run('Q: ').bold = True
        q_para.add_run(f'"{question["question"]}"')
        
        # Add context information
        context_para = doc.add_paragraph()
        context_para.add_run('Context: ').bold = True
        if question.get('document_name'):
            context_para.add_run(f'Document - {question["document_name"]}')
        elif question.get('folder_name'):
            context_para.add_run(f'Folder - {question["folder_name"]}')
        else:
            context_para.add_run('General inquiry')
        
        # Add date asked
        if question.get('created_at'):
            date_asked = datetime.fromisoformat(question['created_at'].replace('Z', '+00:00'))
            context_para.add_run(f' | Asked on {date_asked.strftime("%B %d, %Y at %I:%M %p")}')
        
        # Add answer
        answer_para = doc.add_paragraph()
        answer_para.add_run('A: ').bold = True
        
        if question.get('answer'):
            # Clean the answer text
            clean_answer = clean_markdown_for_word(question['answer'])
            
            # Extract confidence levels for highlighting
            confidence_levels = extract_confidence_level(question['answer'])
            
            # Add the cleaned answer
            answer_para.add_run(clean_answer)
            
            # Add confidence summary if found
            if confidence_levels:
                conf_para = doc.add_paragraph()
                conf_para.add_run('Confidence Indicators: ').bold = True
                conf_para.add_run(', '.join(confidence_levels))
        else:
            no_answer_run = answer_para.add_run('No answer could be determined from the available documentation.')
            no_answer_run.italic = True
        
        # Add referenced documents
        if question.get('referenced_documents') and len(question['referenced_documents']) > 0:
            ref_para = doc.add_paragraph()
            ref_para.add_run('Referenced Documents: ').bold = True
            
            # Group documents by filename to avoid duplicates
            unique_docs = {}
            for doc_ref in question['referenced_documents']:
                filename = doc_ref['filename']
                if filename not in unique_docs:
                    unique_docs[filename] = []
                unique_docs[filename].append(doc_ref['page_number'])
            
            doc_list = []
            for filename, pages in unique_docs.items():
                if len(pages) == 1:
                    doc_list.append(f"{filename} (Page {pages[0]})")
                else:
                    pages_str = ', '.join(map(str, sorted(set(pages))))
                    doc_list.append(f"{filename} (Pages {pages_str})")
            
            ref_para.add_run('; '.join(doc_list))
        
        # Add spacing between questions
        doc.add_paragraph()
    
    # Add summary section at the end
    doc.add_page_break()
    summary_heading = doc.add_heading('Report Summary', level=1)
    
    summary_stats = doc.add_paragraph()
    summary_stats.add_run('Statistics:\n').bold = True
    summary_stats.add_run(f'• Total Questions: {total_questions}\n')
    summary_stats.add_run(f'• Answered Questions: {answered_questions}\n')
    summary_stats.add_run(f'• Unanswered Questions: {unanswered_questions}\n')
    
    if answered_questions > 0:
        answer_rate = (answered_questions / total_questions) * 100
        summary_stats.add_run(f'• Answer Rate: {answer_rate:.1f}%')
    
    return doc

def main(req: func.HttpRequest) -> func.HttpResponse:
    
    if req.headers.get('function-key') != os.environ["FUNCTION_KEY"]:
        logging.info("no matching value in header")
        return func.HttpResponse("", status_code=401)
    
    try:
        email, err = auth_get_email(req)
        if err:
            return err
        
        dd_id = req.route_params.get('dd_id')
        if not dd_id:
            return func.HttpResponse("Missing dd_id parameter", status_code=400)
        
        with transactional_session() as session:
            # Get DD info for title
            dd_info = session.query(DueDiligence).filter(DueDiligence.id == dd_id).first()
            dd_title = dd_info.name if dd_info else "Due Diligence"
            
            # Get all questions for this DD with their referenced documents
            questions = (
                session.query(DDQuestion)
                .options(joinedload(DDQuestion.referenced_documents))
                .filter(DDQuestion.dd_id == dd_id)
                .order_by(DDQuestion.created_at.desc())
                .all()
            )
            
            if not questions:
                return func.HttpResponse("No questions found for this due diligence", status_code=404)
            
            # Convert to dict format
            questions_data = []
            for question in questions:
                question_dict = {
                    "id": str(question.id),
                    "question": question.question,
                    "answer": question.answer,
                    "asked_by": question.asked_by,
                    "folder_id": str(question.folder_id) if question.folder_id else None,
                    "document_id": str(question.document_id) if question.document_id else None,
                    "folder_name": question.folder_name,
                    "document_name": question.document_name,
                    "created_at": question.created_at.isoformat() if question.created_at else None,
                    "referenced_documents": [
                        {
                            "doc_id": str(ref_doc.doc_id),
                            "filename": ref_doc.filename,
                            "page_number": ref_doc.page_number,
                            "folder_path": ref_doc.folder_path
                        }
                        for ref_doc in question.referenced_documents
                    ]
                }
                questions_data.append(question_dict)
            
            # Create Word document
            doc = create_word_document(questions_data, dd_title)
            
            # Save document to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Generate filename
            safe_title = re.sub(r'[^\w\s-]', '', dd_title).strip()
            safe_title = re.sub(r'[-\s]+', '-', safe_title)
            filename = f"DD-QA-Report-{safe_title}-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
            
            return func.HttpResponse(
                doc_buffer.getvalue(),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename={filename}",
                    "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                },
                status_code=200
            )
    
    except Exception as e:
        logging.exception("Error generating Word document")
        return func.HttpResponse(f"Error: {str(e)}", status_code=500)