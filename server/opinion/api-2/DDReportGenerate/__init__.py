"""
API endpoint for generating DD reports and exports.
Supports PDF, DOCX, and Excel formats with blob storage upload.

Phase 7: Enterprise Features
"""

import azure.functions as func
import json
import logging
import os
import uuid
from typing import Dict, List, Any, Optional
from datetime import datetime
from io import BytesIO
from sqlalchemy import text

from shared.session import transactional_session
from shared.uploader import write_to_blob_storage, get_blob_sas_url
from shared.audit import log_audit_event, AuditEventType

logger = logging.getLogger(__name__)

# Environment variables for blob storage
BLOB_CONNECTION_STRING = os.environ.get("DD_DOCS_BLOB_STORAGE_CONNECTION_STRING", "")
BLOB_CONTAINER_NAME = os.environ.get("DD_DOCS_STORAGE_CONTAINER_NAME", "dd-documents")
REPORTS_CONTAINER_NAME = os.environ.get("DD_REPORTS_STORAGE_CONTAINER_NAME", "dd-reports")


def main(req: func.HttpRequest) -> func.HttpResponse:
    """
    Report generation endpoint.

    Routes:
        GET /dd-report-generate?action=list_templates
        GET /dd-report-generate?action=get_export_jobs&dd_id=X
        POST /dd-report-generate (body: action, dd_id, format, template_id, ...)
    """
    try:
        if req.method == 'OPTIONS':
            return func.HttpResponse(status_code=200)

        if req.method == 'GET':
            return handle_get(req)
        elif req.method == 'POST':
            return handle_post(req)
        else:
            return func.HttpResponse(
                json.dumps({"error": f"Method {req.method} not allowed"}),
                status_code=405,
                mimetype="application/json"
            )

    except Exception as e:
        logger.error(f"Error in report generation endpoint: {e}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            status_code=500,
            mimetype="application/json"
        )


def handle_get(req: func.HttpRequest) -> func.HttpResponse:
    """Handle GET requests."""
    action = req.params.get('action')

    if not action:
        return func.HttpResponse(
            json.dumps({"error": "action is required"}),
            status_code=400,
            mimetype="application/json"
        )

    with transactional_session() as session:
        if action == 'list_templates':
            data = list_templates(session)
        elif action == 'get_export_jobs':
            dd_id = req.params.get('dd_id')
            if not dd_id:
                return func.HttpResponse(
                    json.dumps({"error": "dd_id is required"}),
                    status_code=400,
                    mimetype="application/json"
                )
            data = get_export_jobs(session, dd_id)
        elif action == 'get_download_url':
            job_id = req.params.get('job_id')
            if not job_id:
                return func.HttpResponse(
                    json.dumps({"error": "job_id is required"}),
                    status_code=400,
                    mimetype="application/json"
                )
            data = get_download_url(session, job_id)
        else:
            return func.HttpResponse(
                json.dumps({"error": f"Unknown action: {action}"}),
                status_code=400,
                mimetype="application/json"
            )

    return func.HttpResponse(
        json.dumps(data, default=str),
        status_code=200,
        mimetype="application/json"
    )


def handle_post(req: func.HttpRequest) -> func.HttpResponse:
    """Handle POST requests for report generation."""
    try:
        body = req.get_json()
    except ValueError:
        return func.HttpResponse(
            json.dumps({"error": "Invalid JSON body"}),
            status_code=400,
            mimetype="application/json"
        )

    action = body.get('action')

    if not action:
        return func.HttpResponse(
            json.dumps({"error": "action is required"}),
            status_code=400,
            mimetype="application/json"
        )

    with transactional_session() as session:
        if action == 'generate_report':
            result = generate_report(session, body)
        elif action == 'export_findings':
            result = export_findings(session, body)
        elif action == 'export_graph':
            result = export_graph_data(session, body)
        else:
            return func.HttpResponse(
                json.dumps({"error": f"Unknown action: {action}"}),
                status_code=400,
                mimetype="application/json"
            )

        session.commit()

    return func.HttpResponse(
        json.dumps(result, default=str),
        status_code=200,
        mimetype="application/json"
    )


# ============================================================================
# Template Functions
# ============================================================================

def list_templates(session) -> Dict:
    """List available report templates."""
    result = session.execute(
        text("""
            SELECT id, name, description, format, report_type, is_active
            FROM dd_report_template
            WHERE is_active = TRUE
            ORDER BY name
        """)
    )

    templates = [
        {
            'id': str(row.id),
            'name': row.name,
            'description': row.description,
            'format': row.format,
            'report_type': row.report_type
        }
        for row in result.fetchall()
    ]

    return {'templates': templates}


# ============================================================================
# Export Job Functions
# ============================================================================

def get_export_jobs(session, dd_id: str) -> Dict:
    """Get export jobs for a DD project."""
    result = session.execute(
        text("""
            SELECT
                j.id,
                j.dd_id,
                j.export_type,
                j.format,
                j.status,
                j.file_path,
                j.error_message,
                j.created_by,
                j.created_at,
                j.completed_at,
                u.name as created_by_name
            FROM dd_export_job j
            LEFT JOIN users u ON j.created_by = u.id
            WHERE j.dd_id = :dd_id
            ORDER BY j.created_at DESC
            LIMIT 50
        """),
        {'dd_id': dd_id}
    )

    jobs = [
        {
            'id': str(row.id),
            'dd_id': str(row.dd_id),
            'export_type': row.export_type,
            'format': row.format,
            'status': row.status,
            'file_path': row.file_path,
            'error_message': row.error_message,
            'created_by': {
                'id': str(row.created_by) if row.created_by else None,
                'name': row.created_by_name
            },
            'created_at': row.created_at.isoformat() if row.created_at else None,
            'completed_at': row.completed_at.isoformat() if row.completed_at else None
        }
        for row in result.fetchall()
    ]

    return {'jobs': jobs}


def get_download_url(session, job_id: str) -> Dict:
    """Get download URL for a completed export job."""
    result = session.execute(
        text("""
            SELECT id, file_path, status, format
            FROM dd_export_job
            WHERE id = :job_id
        """),
        {'job_id': job_id}
    ).fetchone()

    if not result:
        return {"error": "Job not found"}

    if result.status != 'completed':
        return {"error": f"Job status is {result.status}, not ready for download"}

    if not result.file_path:
        return {"error": "No file path available"}

    try:
        url = get_blob_sas_url(
            connection_string=BLOB_CONNECTION_STRING,
            container_name=REPORTS_CONTAINER_NAME,
            key=result.file_path,
            expiry_minutes=60
        )
        return {"url": url, "format": result.format}
    except Exception as e:
        logger.error(f"Error generating download URL: {e}")
        return {"error": f"Failed to generate download URL: {str(e)}"}


def create_export_job(session, dd_id: str, export_type: str, format: str, created_by: Optional[str] = None) -> str:
    """Create a new export job record."""
    job_id = str(uuid.uuid4())

    session.execute(
        text("""
            INSERT INTO dd_export_job
            (id, dd_id, export_type, format, status, created_by, created_at)
            VALUES (:id, :dd_id, :export_type, :format, :status, :created_by, :created_at)
        """),
        {
            'id': job_id,
            'dd_id': dd_id,
            'export_type': export_type,
            'format': format,
            'status': 'pending',
            'created_by': created_by,
            'created_at': datetime.utcnow()
        }
    )

    return job_id


def update_export_job(session, job_id: str, status: str, file_path: Optional[str] = None, error_message: Optional[str] = None):
    """Update export job status."""
    updates = ["status = :status"]
    params = {'job_id': job_id, 'status': status}

    if file_path:
        updates.append("file_path = :file_path")
        params['file_path'] = file_path

    if error_message:
        updates.append("error_message = :error_message")
        params['error_message'] = error_message

    if status in ['completed', 'failed']:
        updates.append("completed_at = :completed_at")
        params['completed_at'] = datetime.utcnow()

    session.execute(
        text(f"UPDATE dd_export_job SET {', '.join(updates)} WHERE id = :job_id"),
        params
    )


# ============================================================================
# Report Generation Functions
# ============================================================================

def generate_report(session, body: Dict) -> Dict:
    """
    Generate a DD report (PDF or DOCX).

    Body:
        dd_id: DD project ID
        run_id: Analysis run ID (optional, uses latest)
        format: 'pdf' or 'docx'
        template_id: Report template ID (optional)
        sections: List of sections to include (optional)
        created_by: User ID
    """
    dd_id = body.get('dd_id')
    run_id = body.get('run_id')
    format = body.get('format', 'pdf')
    template_id = body.get('template_id')
    sections = body.get('sections')
    created_by = body.get('created_by')

    if not dd_id:
        return {"error": "dd_id is required"}

    if format not in ['pdf', 'docx']:
        return {"error": f"Unsupported format: {format}. Use 'pdf' or 'docx'"}

    # Create export job
    job_id = create_export_job(session, dd_id, 'full_report', format, created_by)

    try:
        update_export_job(session, job_id, 'processing')

        # Get latest run if not specified
        if not run_id:
            result = session.execute(
                text("""
                    SELECT id FROM dd_analysis_run
                    WHERE dd_id = :dd_id
                    ORDER BY created_at DESC LIMIT 1
                """),
                {'dd_id': dd_id}
            ).fetchone()
            run_id = str(result.id) if result else None

        if not run_id:
            update_export_job(session, job_id, 'failed', error_message='No analysis runs found')
            return {"error": "No analysis runs found"}

        # Gather report data
        report_data = gather_report_data(session, dd_id, run_id)

        # Generate report based on format
        if format == 'pdf':
            file_bytes = generate_pdf_report(report_data, sections)
        else:
            file_bytes = generate_docx_report(report_data, sections)

        # Upload to blob storage
        file_name = f"dd_report_{dd_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.{format}"
        file_path = f"{dd_id}/reports/{file_name}"

        write_to_blob_storage(
            connection_string=BLOB_CONNECTION_STRING,
            container_name=REPORTS_CONTAINER_NAME,
            key=file_path,
            blob=file_bytes,
            meta_data={'dd_id': dd_id, 'format': format, 'original_file_name': file_name, 'extension': format},
            overwrite=True
        )

        update_export_job(session, job_id, 'completed', file_path=file_path)

        # Log audit event
        log_audit_event(
            session=session,
            event_type=AuditEventType.REPORT_GENERATED.value,
            entity_type='dd',
            entity_id=dd_id,
            user_id=created_by,
            dd_id=dd_id,
            details={'format': format, 'job_id': job_id, 'run_id': run_id}
        )

        return {
            "success": True,
            "job_id": job_id,
            "file_path": file_path,
            "message": f"Report generated successfully in {format.upper()} format"
        }

    except Exception as e:
        logger.error(f"Error generating report: {e}")
        update_export_job(session, job_id, 'failed', error_message=str(e))
        return {"error": str(e), "job_id": job_id}


def gather_report_data(session, dd_id: str, run_id: str) -> Dict:
    """Gather all data needed for the report."""
    # Get DD info
    dd_info = session.execute(
        text("""
            SELECT id, name, transaction_type, target_name, acquirer_name, deal_value, currency
            FROM dd
            WHERE id = :dd_id
        """),
        {'dd_id': dd_id}
    ).fetchone()

    # Get severity breakdown
    severity_result = session.execute(
        text("""
            SELECT
                action_priority as severity,
                COUNT(*) as count
            FROM perspective_risk_finding
            WHERE run_id = :run_id
            AND status != 'Deleted'
            GROUP BY action_priority
        """),
        {'run_id': run_id}
    ).fetchall()

    severity_breakdown = {row.severity or 'none': row.count for row in severity_result}

    # Get deal blockers
    deal_blockers = session.execute(
        text("""
            SELECT
                id,
                COALESCE(title, LEFT(phrase, 100)) as title,
                phrase as description,
                folder_category,
                COALESCE(risk_category, folder_category) as risk_category,
                action_items as recommendation,
                clause_reference
            FROM perspective_risk_finding
            WHERE run_id = :run_id
            AND (action_priority = 'critical' OR deal_impact = 'deal_blocker')
            AND status != 'Deleted'
            ORDER BY created_at
        """),
        {'run_id': run_id}
    ).fetchall()

    # Get top findings by category
    findings_by_category = session.execute(
        text("""
            SELECT
                COALESCE(risk_category, folder_category, 'General') as category,
                id,
                COALESCE(title, LEFT(phrase, 100)) as title,
                phrase as description,
                action_priority as severity,
                action_items as recommendation,
                clause_reference
            FROM perspective_risk_finding
            WHERE run_id = :run_id
            AND status != 'Deleted'
            ORDER BY
                CASE action_priority
                    WHEN 'critical' THEN 1
                    WHEN 'high' THEN 2
                    WHEN 'medium' THEN 3
                    WHEN 'low' THEN 4
                    ELSE 5
                END
            LIMIT 100
        """),
        {'run_id': run_id}
    ).fetchall()

    # Get CoC summary
    coc_summary = session.execute(
        text("""
            SELECT
                COUNT(*) as total,
                SUM(CASE WHEN has_change_of_control THEN 1 ELSE 0 END) as with_coc,
                SUM(CASE WHEN has_consent_requirement THEN 1 ELSE 0 END) as with_consent
            FROM kg_agreement
            WHERE dd_id = :dd_id
        """),
        {'dd_id': dd_id}
    ).fetchone()

    return {
        'dd_info': {
            'id': str(dd_info.id) if dd_info else None,
            'name': dd_info.name if dd_info else 'Untitled',
            'transaction_type': dd_info.transaction_type if dd_info else None,
            'target_name': dd_info.target_name if dd_info else None,
            'acquirer_name': dd_info.acquirer_name if dd_info else None,
            'deal_value': float(dd_info.deal_value) if dd_info and dd_info.deal_value else None,
            'currency': dd_info.currency if dd_info else None
        },
        'run_id': run_id,
        'generated_at': datetime.utcnow().isoformat(),
        'severity_breakdown': severity_breakdown,
        'deal_blockers': [
            {
                'id': str(row.id),
                'title': row.title,
                'description': row.description[:500] if row.description else None,
                'folder_category': row.folder_category,
                'risk_category': row.risk_category,
                'recommendation': row.recommendation,
                'clause_reference': row.clause_reference
            }
            for row in deal_blockers
        ],
        'findings_by_category': [
            {
                'category': row.category,
                'id': str(row.id),
                'title': row.title,
                'description': row.description[:300] if row.description else None,
                'severity': row.severity,
                'recommendation': row.recommendation,
                'clause_reference': row.clause_reference
            }
            for row in findings_by_category
        ],
        'coc_analysis': {
            'total_agreements': coc_summary.total if coc_summary else 0,
            'agreements_with_coc': coc_summary.with_coc if coc_summary else 0,
            'agreements_requiring_consent': coc_summary.with_consent if coc_summary else 0
        }
    }


def generate_pdf_report(report_data: Dict, sections: Optional[List[str]] = None) -> bytes:
    """Generate PDF report using ReportLab."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    except ImportError:
        raise ImportError("ReportLab is required for PDF generation. Install with: pip install reportlab")

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=20,
        spaceAfter=10
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_JUSTIFY,
        spaceAfter=8
    )
    severity_critical = ParagraphStyle(
        'SeverityCritical',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.red
    )
    severity_high = ParagraphStyle(
        'SeverityHigh',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.orangered
    )

    story = []

    # Title page
    dd_info = report_data.get('dd_info', {})
    story.append(Paragraph("Due Diligence Report", title_style))
    story.append(Spacer(1, 20))

    if dd_info.get('name'):
        story.append(Paragraph(f"<b>Project:</b> {dd_info['name']}", body_style))
    if dd_info.get('transaction_type'):
        story.append(Paragraph(f"<b>Transaction Type:</b> {dd_info['transaction_type']}", body_style))
    if dd_info.get('target_name'):
        story.append(Paragraph(f"<b>Target:</b> {dd_info['target_name']}", body_style))
    if dd_info.get('acquirer_name'):
        story.append(Paragraph(f"<b>Acquirer:</b> {dd_info['acquirer_name']}", body_style))
    if dd_info.get('deal_value'):
        story.append(Paragraph(
            f"<b>Deal Value:</b> {dd_info.get('currency', 'USD')} {dd_info['deal_value']:,.0f}",
            body_style
        ))

    story.append(Spacer(1, 10))
    story.append(Paragraph(f"<b>Generated:</b> {report_data.get('generated_at', '')[:10]}", body_style))
    story.append(PageBreak())

    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))

    severity = report_data.get('severity_breakdown', {})
    total_findings = sum(severity.values())
    critical = severity.get('critical', 0)
    high = severity.get('high', 0)
    medium = severity.get('medium', 0)
    low = severity.get('low', 0)

    story.append(Paragraph(
        f"This due diligence review identified <b>{total_findings}</b> findings across the document set.",
        body_style
    ))

    # Severity summary table
    severity_data = [
        ['Severity', 'Count'],
        ['Critical', str(critical)],
        ['High', str(high)],
        ['Medium', str(medium)],
        ['Low', str(low)]
    ]
    severity_table = Table(severity_data, colWidths=[3*inch, 1.5*inch])
    severity_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F46E5')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#FEE2E2')),  # Critical row
        ('BACKGROUND', (0, 2), (-1, 2), colors.HexColor('#FED7AA')),  # High row
        ('BACKGROUND', (0, 3), (-1, 3), colors.HexColor('#FEF3C7')),  # Medium row
        ('BACKGROUND', (0, 4), (-1, 4), colors.HexColor('#D1FAE5')),  # Low row
    ]))
    story.append(severity_table)
    story.append(Spacer(1, 20))

    # Deal blockers section
    deal_blockers = report_data.get('deal_blockers', [])
    if deal_blockers:
        story.append(Paragraph("Deal Blockers", heading_style))
        story.append(Paragraph(
            f"<b>{len(deal_blockers)}</b> critical issues have been identified that may impact the transaction:",
            body_style
        ))
        story.append(Spacer(1, 10))

        for i, blocker in enumerate(deal_blockers, 1):
            story.append(Paragraph(f"<b>{i}. {blocker.get('title', 'Untitled')}</b>", severity_critical))
            if blocker.get('description'):
                story.append(Paragraph(blocker['description'][:400], body_style))
            if blocker.get('recommendation'):
                story.append(Paragraph(f"<i>Recommendation: {blocker['recommendation']}</i>", body_style))
            story.append(Spacer(1, 10))

    # CoC Analysis
    coc = report_data.get('coc_analysis', {})
    if coc.get('total_agreements', 0) > 0:
        story.append(Paragraph("Change of Control Analysis", heading_style))
        story.append(Paragraph(
            f"Of <b>{coc['total_agreements']}</b> agreements analyzed:",
            body_style
        ))
        story.append(Paragraph(
            f"- <b>{coc.get('agreements_with_coc', 0)}</b> contain change of control provisions",
            body_style
        ))
        story.append(Paragraph(
            f"- <b>{coc.get('agreements_requiring_consent', 0)}</b> require consent for assignment/transfer",
            body_style
        ))

    # Findings by category
    findings = report_data.get('findings_by_category', [])
    if findings:
        story.append(PageBreak())
        story.append(Paragraph("Findings by Category", heading_style))

        # Group by category
        categories = {}
        for f in findings:
            cat = f.get('category', 'General')
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(f)

        for category, cat_findings in categories.items():
            story.append(Paragraph(f"<b>{category}</b> ({len(cat_findings)} findings)", body_style))
            story.append(Spacer(1, 5))

            for finding in cat_findings[:5]:  # Limit to 5 per category
                severity_label = finding.get('severity', 'low').upper()
                story.append(Paragraph(
                    f"[{severity_label}] {finding.get('title', 'Untitled')}",
                    severity_high if severity_label in ['CRITICAL', 'HIGH'] else body_style
                ))

            story.append(Spacer(1, 10))

    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx_report(report_data: Dict, sections: Optional[List[str]] = None) -> bytes:
    """Generate DOCX report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")

    doc = Document()

    # Title
    title = doc.add_heading('Due Diligence Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project info
    dd_info = report_data.get('dd_info', {})

    if dd_info.get('name'):
        doc.add_paragraph(f"Project: {dd_info['name']}")
    if dd_info.get('transaction_type'):
        doc.add_paragraph(f"Transaction Type: {dd_info['transaction_type']}")
    if dd_info.get('target_name'):
        doc.add_paragraph(f"Target: {dd_info['target_name']}")
    if dd_info.get('acquirer_name'):
        doc.add_paragraph(f"Acquirer: {dd_info['acquirer_name']}")
    if dd_info.get('deal_value'):
        doc.add_paragraph(f"Deal Value: {dd_info.get('currency', 'USD')} {dd_info['deal_value']:,.0f}")

    doc.add_paragraph(f"Generated: {report_data.get('generated_at', '')[:10]}")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)

    severity = report_data.get('severity_breakdown', {})
    total_findings = sum(severity.values())

    doc.add_paragraph(
        f"This due diligence review identified {total_findings} findings across the document set."
    )

    # Severity table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['Severity', 'Count']
    rows_data = [
        ('Critical', severity.get('critical', 0)),
        ('High', severity.get('high', 0)),
        ('Medium', severity.get('medium', 0)),
        ('Low', severity.get('low', 0))
    ]

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = headers[0]
    hdr_cells[1].text = headers[1]

    for i, (sev, count) in enumerate(rows_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = sev
        row_cells[1].text = str(count)

    doc.add_paragraph()

    # Deal blockers
    deal_blockers = report_data.get('deal_blockers', [])
    if deal_blockers:
        doc.add_heading('Deal Blockers', level=1)
        doc.add_paragraph(f"{len(deal_blockers)} critical issues have been identified:")

        for i, blocker in enumerate(deal_blockers, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {blocker.get('title', 'Untitled')}")
            run.bold = True

            if blocker.get('description'):
                doc.add_paragraph(blocker['description'][:400])

            if blocker.get('recommendation'):
                p = doc.add_paragraph()
                run = p.add_run(f"Recommendation: {blocker['recommendation']}")
                run.italic = True

    # CoC Analysis
    coc = report_data.get('coc_analysis', {})
    if coc.get('total_agreements', 0) > 0:
        doc.add_heading('Change of Control Analysis', level=1)
        doc.add_paragraph(f"Of {coc['total_agreements']} agreements analyzed:")
        doc.add_paragraph(f"- {coc.get('agreements_with_coc', 0)} contain change of control provisions")
        doc.add_paragraph(f"- {coc.get('agreements_requiring_consent', 0)} require consent for assignment/transfer")

    # Findings by category
    findings = report_data.get('findings_by_category', [])
    if findings:
        doc.add_page_break()
        doc.add_heading('Findings by Category', level=1)

        # Group by category
        categories = {}
        for f in findings:
            cat = f.get('category', 'General')
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(f)

        for category, cat_findings in categories.items():
            doc.add_heading(f"{category} ({len(cat_findings)} findings)", level=2)

            for finding in cat_findings[:5]:
                severity_label = finding.get('severity', 'low').upper()
                p = doc.add_paragraph()
                run = p.add_run(f"[{severity_label}] ")
                if severity_label in ['CRITICAL', 'HIGH']:
                    run.font.color.rgb = RGBColor(220, 38, 38)
                p.add_run(finding.get('title', 'Untitled'))

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
# Excel Export Functions
# ============================================================================

def export_findings(session, body: Dict) -> Dict:
    """
    Export findings to Excel.

    Body:
        dd_id: DD project ID
        run_id: Analysis run ID (optional)
        created_by: User ID
    """
    dd_id = body.get('dd_id')
    run_id = body.get('run_id')
    created_by = body.get('created_by')

    if not dd_id:
        return {"error": "dd_id is required"}

    job_id = create_export_job(session, dd_id, 'findings_export', 'xlsx', created_by)

    try:
        update_export_job(session, job_id, 'processing')

        # Get latest run if not specified
        if not run_id:
            result = session.execute(
                text("""
                    SELECT id FROM dd_analysis_run
                    WHERE dd_id = :dd_id
                    ORDER BY created_at DESC LIMIT 1
                """),
                {'dd_id': dd_id}
            ).fetchone()
            run_id = str(result.id) if result else None

        if not run_id:
            update_export_job(session, job_id, 'failed', error_message='No analysis runs found')
            return {"error": "No analysis runs found"}

        # Get findings
        findings = session.execute(
            text("""
                SELECT
                    id,
                    COALESCE(title, LEFT(phrase, 100)) as title,
                    phrase as description,
                    action_priority as severity,
                    COALESCE(risk_category, folder_category) as risk_category,
                    folder_category,
                    action_items as recommendation,
                    clause_reference,
                    deal_impact,
                    financial_exposure_amount,
                    financial_exposure_currency,
                    status,
                    created_at
                FROM perspective_risk_finding
                WHERE run_id = :run_id
                AND status != 'Deleted'
                ORDER BY
                    CASE action_priority
                        WHEN 'critical' THEN 1
                        WHEN 'high' THEN 2
                        WHEN 'medium' THEN 3
                        WHEN 'low' THEN 4
                        ELSE 5
                    END,
                    created_at
            """),
            {'run_id': run_id}
        ).fetchall()

        # Generate Excel
        file_bytes = generate_findings_excel(findings)

        # Upload
        file_name = f"dd_findings_{dd_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.xlsx"
        file_path = f"{dd_id}/exports/{file_name}"

        write_to_blob_storage(
            connection_string=BLOB_CONNECTION_STRING,
            container_name=REPORTS_CONTAINER_NAME,
            key=file_path,
            blob=file_bytes,
            meta_data={'dd_id': dd_id, 'format': 'xlsx', 'original_file_name': file_name, 'extension': 'xlsx'},
            overwrite=True
        )

        update_export_job(session, job_id, 'completed', file_path=file_path)

        log_audit_event(
            session=session,
            event_type=AuditEventType.REPORT_EXPORTED.value,
            entity_type='dd',
            entity_id=dd_id,
            user_id=created_by,
            dd_id=dd_id,
            details={'format': 'xlsx', 'job_id': job_id, 'findings_count': len(findings)}
        )

        return {
            "success": True,
            "job_id": job_id,
            "file_path": file_path,
            "findings_count": len(findings)
        }

    except Exception as e:
        logger.error(f"Error exporting findings: {e}")
        update_export_job(session, job_id, 'failed', error_message=str(e))
        return {"error": str(e), "job_id": job_id}


def generate_findings_excel(findings) -> bytes:
    """Generate Excel file from findings."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise ImportError("openpyxl is required for Excel generation. Install with: pip install openpyxl")

    wb = Workbook()
    ws = wb.active
    ws.title = "Findings"

    # Headers
    headers = [
        'ID', 'Title', 'Description', 'Severity', 'Risk Category',
        'Folder Category', 'Recommendation', 'Clause Reference',
        'Deal Impact', 'Financial Exposure', 'Currency', 'Status', 'Created At'
    ]

    # Style definitions
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='4F46E5', end_color='4F46E5', fill_type='solid')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Write headers
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border

    # Severity colors
    severity_colors = {
        'critical': 'FEE2E2',
        'high': 'FED7AA',
        'medium': 'FEF3C7',
        'low': 'D1FAE5'
    }

    # Write data
    for row_num, finding in enumerate(findings, 2):
        row_data = [
            str(finding.id),
            finding.title,
            (finding.description[:500] if finding.description else ''),
            finding.severity or 'none',
            finding.risk_category,
            finding.folder_category,
            finding.recommendation,
            finding.clause_reference,
            finding.deal_impact,
            float(finding.financial_exposure_amount) if finding.financial_exposure_amount else None,
            finding.financial_exposure_currency,
            finding.status,
            finding.created_at.isoformat() if finding.created_at else ''
        ]

        severity = finding.severity or 'none'
        row_fill = PatternFill(
            start_color=severity_colors.get(severity, 'FFFFFF'),
            end_color=severity_colors.get(severity, 'FFFFFF'),
            fill_type='solid'
        )

        for col, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col, value=value)
            cell.border = thin_border
            cell.fill = row_fill

    # Auto-adjust column widths
    for col in range(1, len(headers) + 1):
        max_length = len(headers[col - 1])
        for row in range(2, min(len(findings) + 2, 100)):  # Check first 100 rows
            cell_value = ws.cell(row=row, column=col).value
            if cell_value:
                max_length = max(max_length, min(len(str(cell_value)), 50))
        ws.column_dimensions[get_column_letter(col)].width = max_length + 2

    # Save to buffer
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_graph_data(session, body: Dict) -> Dict:
    """
    Export knowledge graph data to JSON.

    Body:
        dd_id: DD project ID
        created_by: User ID
    """
    dd_id = body.get('dd_id')
    created_by = body.get('created_by')

    if not dd_id:
        return {"error": "dd_id is required"}

    job_id = create_export_job(session, dd_id, 'graph_export', 'json', created_by)

    try:
        update_export_job(session, job_id, 'processing')

        # Gather graph data
        graph_data = {
            'dd_id': dd_id,
            'exported_at': datetime.utcnow().isoformat(),
            'parties': [],
            'agreements': [],
            'triggers': [],
            'obligations': [],
            'edges': []
        }

        # Get parties
        parties = session.execute(
            text("""
                SELECT id, name, party_type, role, jurisdiction
                FROM kg_party WHERE dd_id = :dd_id
            """),
            {'dd_id': dd_id}
        ).fetchall()

        graph_data['parties'] = [
            {
                'id': str(p.id), 'name': p.name, 'party_type': p.party_type,
                'role': p.role, 'jurisdiction': p.jurisdiction
            }
            for p in parties
        ]

        # Get agreements
        agreements = session.execute(
            text("""
                SELECT id, name, agreement_type, has_change_of_control,
                       has_consent_requirement, has_assignment_restriction
                FROM kg_agreement WHERE dd_id = :dd_id
            """),
            {'dd_id': dd_id}
        ).fetchall()

        graph_data['agreements'] = [
            {
                'id': str(a.id), 'name': a.name, 'agreement_type': a.agreement_type,
                'has_coc': a.has_change_of_control, 'has_consent': a.has_consent_requirement,
                'has_assignment_restriction': a.has_assignment_restriction
            }
            for a in agreements
        ]

        # Get triggers
        triggers = session.execute(
            text("""
                SELECT id, trigger_type, description, consequences, agreement_id
                FROM kg_trigger WHERE dd_id = :dd_id
            """),
            {'dd_id': dd_id}
        ).fetchall()

        graph_data['triggers'] = [
            {
                'id': str(t.id), 'trigger_type': t.trigger_type,
                'description': t.description, 'consequences': t.consequences,
                'agreement_id': str(t.agreement_id) if t.agreement_id else None
            }
            for t in triggers
        ]

        # Export as JSON
        json_bytes = json.dumps(graph_data, indent=2, default=str).encode('utf-8')

        file_name = f"dd_graph_{dd_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json"
        file_path = f"{dd_id}/exports/{file_name}"

        write_to_blob_storage(
            connection_string=BLOB_CONNECTION_STRING,
            container_name=REPORTS_CONTAINER_NAME,
            key=file_path,
            blob=json_bytes,
            meta_data={'dd_id': dd_id, 'format': 'json', 'original_file_name': file_name, 'extension': 'json'},
            overwrite=True
        )

        update_export_job(session, job_id, 'completed', file_path=file_path)

        return {
            "success": True,
            "job_id": job_id,
            "file_path": file_path,
            "stats": {
                'parties': len(graph_data['parties']),
                'agreements': len(graph_data['agreements']),
                'triggers': len(graph_data['triggers'])
            }
        }

    except Exception as e:
        logger.error(f"Error exporting graph data: {e}")
        update_export_job(session, job_id, 'failed', error_message=str(e))
        return {"error": str(e), "job_id": job_id}
