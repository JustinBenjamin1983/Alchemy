import logging, os, json, io, re
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Dict, Any, List, Tuple
import azure.functions as func
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate, RichText, RichTextParagraph
from shared.utils import auth_get_email, generate_identifier
from shared.table_storage import get_user_info
from shared.rag import call_llm_with
from shared.uploader import read_from_blob_storage, write_to_blob_storage, get_blob_sas_url

# This is a test

TZ_SA = timezone(timedelta(hours=2))

CHUNK_SIZE = 50000

def extract_and_number_citations(opinion_text: str) -> dict:
    
    system_prompt = """You are a legal citation specialist for South African law firms.

Your task is to:
1. Extract ALL citations from the legal opinion (statutes, cases, documents)
2. Assign each unique citation a number [1], [2], [3], etc.
3. Replace citation text with numbered markers
4. Create a structured reference list

CITATION TYPES TO EXTRACT:
1. **Statutes/Acts**: Companies Act, No. 71 of 2008, PAIA, etc.
2. **Case Law**: Full case names with citations and SAFLII URLs
3. **Client Documents**: Agreement of Sale, contracts, etc.
4. **Regulations**: JSE Listings Requirements, etc.

RULES:
- Each unique source gets ONE number
- First mention should be the full citation followed by [N]
- Subsequent mentions can use short form followed by [N]
- SAFLII cases MUST include the full URL in references
- Document citations should include page numbers in the reference list

OUTPUT FORMAT (strict JSON):
{
  "numbered_opinion": "Full opinion text with [1], [2], etc. markers inserted",
  "references": [
    {
      "number": 1,
      "type": "statute|case|document|regulation",
      "full_citation": "Complete formal citation",
      "short_citation": "Brief form for subsequent use",
      "url": "https://... (for SAFLII cases)",
      "pages": "Page references (for documents)"
    }
  ]
}

EXAMPLE INPUT:
"The Companies Act, No. 71 of 2008 provides that... As held in Minister of Mineral Resources v Sustaining the Wild Coast [2024] ZASCA 84..."

EXAMPLE OUTPUT:
{
  "numbered_opinion": "The Companies Act, No. 71 of 2008 [1] provides that... As held in Minister of Mineral Resources v Sustaining the Wild Coast [2024] ZASCA 84 [2]...",
  "references": [
    {
      "number": 1,
      "type": "statute",
      "full_citation": "Companies Act, No. 71 of 2008",
      "short_citation": "Companies Act",
      "url": null,
      "pages": null
    },
    {
      "number": 2,
      "type": "case",
      "full_citation": "Minister of Mineral Resources and Energy and Others v Sustaining the Wild Coast NPC and Others [2024] ZASCA 84; 2024 (5) SA 38 (SCA)",
      "short_citation": "Sustaining the Wild Coast",
      "url": "https://www.saflii.org/za/cases/ZASCA/2024/84.html",
      "pages": null
    }
  ]
}

CRITICAL:
- Return ONLY valid JSON
- Number citations sequentially in order of first appearance
- Preserve all original text, only add [N] markers
- Include full SAFLII URLs for all cases
"""

    user_prompt = f"""Please extract all citations from this legal opinion and add numbered references:

OPINION TEXT:
{opinion_text}

Return the numbered opinion and complete reference list in JSON format."""

    try:
        response = call_llm_with(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.0,
            max_tokens=16000,  # Large opinions need space
            model_deployment_env_var='OPINION_MODEL_DEPLOYMENT',
            model_version_env_var='OPINION_MODEL_VERSION'
        )
        
        # Parse JSON response
        try:
            start = response.find('{')
            end = response.rfind('}') + 1
            if start != -1 and end > start:
                json_str = response[start:end]
                result = json.loads(json_str)
            else:
                result = json.loads(response)
            
            logging.info(f"📚 Extracted {len(result.get('references', []))} citations")
            return result
            
        except json.JSONDecodeError as e:
            logging.error(f"Failed to parse citation extraction JSON: {e}")
            logging.error(f"Response was: {response[:1000]}")
            # Return original text without numbering
            return {
                "numbered_opinion": opinion_text,
                "references": []
            }
            
    except Exception as e:
        logging.exception(f"Error in extract_and_number_citations: {e}")
        return {
            "numbered_opinion": opinion_text,
            "references": []
        }


def format_references_section(references: list) -> str:
    """
    Formats the references list into a proper References section for the opinion.
    
    Groups by type: Legislation → Case Law → Documents
    """
    
    if not references:
        return ""
    
    # Group references by type
    statutes = []
    cases = []
    documents = []
    regulations = []
    
    for ref in references:
        ref_type = ref.get("type", "document")
        if ref_type == "statute":
            statutes.append(ref)
        elif ref_type == "case":
            cases.append(ref)
        elif ref_type == "regulation":
            regulations.append(ref)
        else:
            documents.append(ref)
    
    # Build the references section
    sections = []
    
    sections.append("\n\n---\n## REFERENCES\n")
    
    if statutes:
        sections.append("\n### Legislation and Statutes\n")
        for ref in statutes:
            line = f"[{ref['number']}] {ref['full_citation']}"
            sections.append(line)
    
    if regulations:
        sections.append("\n### Regulations and Rules\n")
        for ref in regulations:
            line = f"[{ref['number']}] {ref['full_citation']}"
            sections.append(line)
    
    if cases:
        sections.append("\n### Case Law\n")
        for ref in cases:
            line = f"[{ref['number']}] {ref['full_citation']}"
            if ref.get("url"):
                line += f"\n    URL: {ref['url']}"
            sections.append(line)
    
    if documents:
        sections.append("\n### Client Documents and Materials\n")
        for ref in documents:
            line = f"[{ref['number']}] {ref['full_citation']}"
            if ref.get("pages"):
                line += f" (Pages: {ref['pages']})"
            sections.append(line)
    
    return "\n".join(sections)


def add_citations_to_opinion(opinion_text: str) -> str:
    """
    Main function to add numbered citations to an opinion.
    
    1. Extracts all citations using LLM
    2. Adds [1], [2] markers throughout text
    3. Appends formatted References section
    
    Returns: Complete opinion with numbered citations and references
    """
    
    logging.info("📚 Starting citation extraction and numbering...")
    
    # Extract and number citations
    citation_data = extract_and_number_citations(opinion_text)
    
    numbered_opinion = citation_data.get("numbered_opinion", opinion_text)
    references = citation_data.get("references", [])
    
    if not references:
        logging.warning("⚠️ No citations were extracted - returning original opinion")
        return opinion_text
    
    # Format references section
    references_section = format_references_section(references)
    
    # Append references to opinion (before Audit Appendix if it exists)
    if "## Audit Appendix" in numbered_opinion:
        # Insert before Audit Appendix
        parts = numbered_opinion.split("## Audit Appendix")
        final_opinion = parts[0] + references_section + "\n\n## Audit Appendix" + parts[1]
    else:
        # Append at end
        final_opinion = numbered_opinion + references_section
    
    logging.info(f"✅ Added {len(references)} numbered citations to opinion")
    
    return final_opinion

def _conn_str() -> str:
    cs = (
        os.environ.get("BLOB_STORAGE_CONNECTION_STRING")
        or os.environ.get("USER_TABLE_STORAGE_CONNECTION_STRING")
        or os.environ.get("AZURE_STORAGE_CONNECTION_STRING")
    )
    if not cs:
        raise RuntimeError("No blob storage connection string found in env")
    return cs

def convert_lawyer_opinion_to_client_version(opinion_text: str, client_name: str, subject: str, 
                                                    facts: str, questions: str, assumptions: str) -> dict:
    """Convert lawyer's technical opinion into comprehensive client-friendly version with flexible structure"""
    
    system_prompt = f"""You are a senior partner at a top-tier South African law firm, writing a comprehensive legal opinion specifically for your business client.

CRITICAL TRANSFORMATION REQUIREMENTS:
1. AUDIENCE: This is for {client_name} (business executives), NOT lawyers
2. PURPOSE: Transform the technical legal opinion into an extensive, business-focused advisory document
3. LENGTH: Make this comprehensive and detailed - aim for 3000-5000 words
4. TONE: Professional but accessible business communication

CONTENT TRANSFORMATION RULES:

**PRESERVE (MANDATORY):**
- Numbered citations [1], [2], [3] - keep these EXACTLY as they appear
- The References section at the end - keep this intact with the heading "REFERENCES"
- All factual citations and source references
- Technical accuracy of legal conclusions

**REMOVE:**
- Technical legal jargon and Latin phrases (but keep citations!)
- Lawyer-to-lawyer technical analysis
- "Sources" and "Audit Appendix" sections (but NOT the References section)
- Overly formal language that obscures meaning

**ADD:**
- Detailed explanations of legal concepts in business terms
- Practical business implications of each point
- Clear recommendations with implementation steps
- Risk analysis in business language where relevant
- Timeline and next steps guidance where appropriate
- Cost implications if relevant to the issues

**FLEXIBLE STRUCTURE GUIDELINES:**

You have creative freedom to organize the opinion in the way that best serves this specific client and matter. Choose section names and organization that make sense for THIS case.

**Suggested Section Types** (use what's appropriate, add others as needed):
- Opening summary/overview (choose appropriate name)
- Background/context/situation explanation
- Legal analysis sections (name based on actual issues)
- Business implications/impact sections
- Risk assessment (if significant risks exist)
- Options analysis (if multiple paths available)
- Recommendations (always include)
- Implementation/next steps
- Timeline (if time-sensitive)
- Cost considerations (if financial impact is significant)
- Regulatory compliance (if relevant)
- REFERENCES (mandatory final section)

**Section Naming Guidelines:**
- Use clear, business-focused names like:
  ✓ "What This Means For Your Business"
  ✓ "Your Options and Their Implications"
  ✓ "Understanding Your Rights in This Situation"
  ✓ "Key Risks and How to Manage Them"
  ✓ "The Path Forward: Our Recommendations"
  
- Avoid generic lawyer-speak like:
  ✗ "Opinion and Legal Position"
  ✗ "Issues for Determination"
  ✗ "Legal Analysis of Questions Posed"

**JSON Output Structure:**

{{
  "body_sections": [
    {{
      "type": "heading",
      "text": "[Your chosen section name - make it relevant and clear]",
      "style": {{"bold": true, "size": 26}}
    }},
    {{
      "type": "paragraph",
      "text": "[Content with citations preserved [1], [2]]",
      "style": {{"size": 22}}
    }},
    {{
      "type": "heading",
      "text": "[Another section name - choose what makes sense]",
      "style": {{"bold": true, "size": 26}}
    }},
    {{
      "type": "paragraph",
      "text": "[More content with citations [3], [4]]",
      "style": {{"size": 22}}
    }},
    {{
      "type": "list_item",
      "text": "• Bullet point for recommendations or key points [5]",
      "style": {{"size": 22}}
    }},
    {{
      "type": "heading",
      "text": "REFERENCES",
      "style": {{"bold": true, "size": 24}}
    }},
    {{
      "type": "paragraph",
      "text": "[Preserve entire REFERENCES section exactly as formatted]",
      "style": {{"size": 20}}
    }}
  ]
}}

**Available Section Types:**
- "heading" - For major section titles (size 26, bold)
- "paragraph" - For body text (size 22)
- "list_item" - For bullet points (size 22, use "• " prefix)

**Content Organization Principles:**
1. **Flow Naturally**: Organize content in the order that makes most sense for understanding
2. **Group Related Ideas**: Combine related legal points under meaningful section headings
3. **Emphasize What Matters**: Give more space to issues that have bigger business impact
4. **Be Practical**: Always connect legal analysis to business reality
5. **End Strong**: Finish with clear action items (before References section)

**CRITICAL CITATION RULES:**
- PRESERVE all [1], [2], [3] numbered citations exactly where they appear
- DO NOT renumber or remove citations
- PRESERVE the entire REFERENCES section at the end with the exact heading "REFERENCES"
- Citations prove credibility - keep them visible but explain the content in plain language

**WRITING STYLE:**
- Address the client directly ("your company", "you should", "we recommend")
- Explain WHY things matter from a business perspective
- Use analogies and examples where helpful
- Break complex legal concepts into digestible chunks
- Keep numbered citations to show sources backing your advice
- Use subheadings within sections if it helps clarity
- Vary paragraph length for readability

**CONTEXT FOR THIS OPINION:**
Client: {client_name}
Subject: {subject}
Facts: {facts}
Questions: {questions}
Assumptions: {assumptions}

**YOUR TASK:**
Create a comprehensive, business-focused legal opinion that:
1. Fully explains the legal position AND its business implications
2. Preserves all numbered citations and the References section
3. Uses section names and organization that best serve THIS specific client and matter
4. Makes the client feel informed, advised, and confident about next steps

**IMPORTANT:** 
- Use single spaces between words and clean formatting
- Return ONLY valid JSON
- Be creative with section names but maintain professionalism
- The References section must be the final section before any optional appendices
"""

    user_prompt = f"""Please transform this lawyer's technical opinion into a comprehensive client-friendly business advisory document:

ORIGINAL LAWYER OPINION:
{opinion_text}

Transform this into an extensive, detailed client opinion. Use your judgment to:
- Choose section names that best fit this specific case
- Organize content in the most logical and helpful way for the client
- Add or remove sections as appropriate to the issues involved
- Make it comprehensive and business-focused

Remember to preserve all citations [1], [2], etc. and the References section exactly."""

    try:
        response = call_llm_with(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,  # Slightly higher for more creative section naming
            max_tokens=12000,
            model_deployment_env_var='OPINION_MODEL_DEPLOYMENT',
            model_version_env_var='OPINION_MODEL_VERSION'
        )
        
        # Parse JSON response
        try:
            start = response.find('{')
            end = response.rfind('}') + 1
            if start != -1 and end > start:
                json_str = response[start:end]
                result = json.loads(json_str)
            else:
                result = json.loads(response)
            
            # Log the section structure for debugging
            sections = result.get('body_sections', [])
            section_names = [s.get('text', '') for s in sections if s.get('type') == 'heading']
            logging.info(f"📋 Generated {len(sections)} sections with headings: {section_names}")
            
            return result
            
        except json.JSONDecodeError:
            logging.error(f"Failed to parse client opinion JSON: {response[:500]}")
            return create_client_fallback_structure(opinion_text, client_name, facts, questions)
            
    except Exception as e:
        logging.error(f"Error in convert_lawyer_opinion_to_client_version: {e}")
        return create_client_fallback_structure(opinion_text, client_name, facts, questions)

def create_client_fallback_structure(opinion_text: str, client_name: str, facts: str, questions: str) -> dict:
    """Create a client-focused fallback structure"""
    
    # Clean up the original opinion
    cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'\1', opinion_text)
    cleaned_text = re.sub(r'\[(\d+)\]', '', cleaned_text)
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    return {
        "body_sections": [
            {
                "type": "heading",
                "text": "EXECUTIVE SUMMARY",
                "style": {"bold": True, "size": 26}
            },
            {
                "type": "paragraph",
                "text": f"Dear {client_name}, this opinion addresses your questions about {questions}. Based on our analysis of your situation, we provide our recommendations and next steps below.",
                "style": {"size": 22}
            },
            {
                "type": "heading",
                "text": "YOUR SITUATION", 
                "style": {"bold": True, "size": 26}
            },
            {
                "type": "paragraph",
                "text": f"You have asked us to advise on the following situation: {facts}",
                "style": {"size": 22}
            },
            {
                "type": "heading",
                "text": "OUR ANALYSIS",
                "style": {"bold": True, "size": 26}
            },
            {
                "type": "paragraph",
                "text": cleaned_text,
                "style": {"size": 22}
            },
            {
                "type": "heading",
                "text": "NEXT STEPS",
                "style": {"bold": True, "size": 26}
            },
            {
                "type": "paragraph",
                "text": "We recommend discussing these findings with your management team and implementing the suggestions outlined above. Please contact us if you need clarification on any points.",
                "style": {"size": 22}
            }
        ]
    }
    
def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> List[str]:
    """Split text into chunks for processing."""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at paragraph boundary
        if end < len(text):
            # Look for double newline
            break_point = text.rfind('\n\n', start, end)
            if break_point == -1:
                break_point = text.rfind('\n', start, end)
            if break_point == -1:
                break_point = text.rfind('. ', start, end)
            
            if break_point > start:
                end = break_point + 1
        
        chunks.append(text[start:end])
        start = end
    
    return chunks

def _extract_structured_content_from_chunk(chunk: str, chunk_index: int) -> List[Dict[str, Any]]:
    """
    Extract structured elements (headings, paragraphs, lists) from a text chunk.
    """
    logging.info(f"  🔍 Processing chunk {chunk_index} ({len(chunk):,} chars)")
    
    system_prompt = """You are a legal document structure analyzer. Extract ALL content from the text, identifying its structural type.

**YOUR TASK:**
Analyze the legal opinion text and extract EVERY piece of content, classifying each as:
- **heading_1**: Main section titles (e.g., "EXECUTIVE SUMMARY", "OPINION AND LEGAL POSITION")
- **heading_2**: Subsection titles (e.g., "Issue 1: ...", "Standard of Review")
- **heading_3**: Sub-subsection titles (e.g., "4.1 Conclusion", "Counter-arguments")
- **paragraph**: Regular body text paragraphs
- **list_item**: Bulleted or numbered list items
- **citation_block**: Reference lists or citation sections

**CRITICAL RULES:**
1. Extract ALL content - do not skip anything
2. Do not summarize - preserve exact text including [1], [2] citations
3. Maintain original order
4. Preserve numbered citations [1], [2], [3] exactly as they appear
5. If a section starts with "##" it's likely a heading
6. Lists starting with "-", "•", or numbers are list_items
7. The "REFERENCES" section should be citation_block type

**OUTPUT FORMAT (JSON):**
{
  "elements": [
    {
      "type": "heading_1",
      "content": "EXECUTIVE SUMMARY",
      "order": 0
    },
    {
      "type": "paragraph",
      "content": "This opinion addresses whether [1] the transaction...",
      "order": 1
    },
    {
      "type": "list_item",
      "content": "The Companies Act, No. 71 of 2008 [1] provides that...",
      "order": 2
    }
  ]
}

Respond with ONLY valid JSON containing ALL content from the text."""

    try:
        response = call_llm_with(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Extract ALL structured content from this legal opinion text:\n\n{chunk}"}
            ],
            temperature=0,
            max_tokens=12000,
            model_deployment_env_var='OPINION_MODEL_DEPLOYMENT',
            model_version_env_var='OPINION_MODEL_VERSION'
        )
        
        # Parse JSON response
        result_text = response.strip()
        
        # Try to extract JSON from response
        start = result_text.find('{')
        end = result_text.rfind('}') + 1
        if start != -1 and end > start:
            json_str = result_text[start:end]
            result = json.loads(json_str)
        else:
            result = json.loads(result_text)
        
        elements = result.get("elements", [])
        
        # Add metadata
        for element in elements:
            element["chunk_index"] = chunk_index
        
        logging.info(f"    ✅ Extracted {len(elements)} elements")
        
        return elements
    
    except Exception as e:
        logging.error(f"    ❌ Failed to extract from chunk: {e}")
        # Fallback: treat entire chunk as paragraph
        return [{
            "type": "paragraph",
            "content": chunk,
            "order": 0,
            "chunk_index": chunk_index
        }]

def _process_opinion_text(opinion_text: str, working_dir: Path) -> List[Dict[str, Any]]:
    """
    Process opinion text: chunk it, extract structure.
    Returns list of structured elements.
    """
    logging.info(f"📄 Processing opinion text ({len(opinion_text):,} chars)")
    
    # Chunk the content
    chunks = _chunk_text(opinion_text)
    logging.info(f"  📦 Split into {len(chunks)} chunks")
    
    # Extract structure from each chunk
    all_elements = []
    
    for chunk_idx, chunk in enumerate(chunks):
        elements = _extract_structured_content_from_chunk(chunk, chunk_idx)
        all_elements.extend(elements)
    
    # Save structured content for debugging
    output_file = working_dir / "opinion_structured.json"
    
    structured_data = {
        "total_elements": len(all_elements),
        "total_chunks": len(chunks),
        "original_size": len(opinion_text),
        "elements": all_elements
    }
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(structured_data, f, indent=2, ensure_ascii=False)
    
    logging.info(f"  ✅ Extracted {len(all_elements)} structured elements")
    
    return all_elements

def _is_custom_style(style_name: str) -> bool:
    """
    Check if a style name is a custom style based on specific prefixes.
    Custom styles must start with: Alchemy, Annexure, or Schedule (case-insensitive).
    """
    style_lower = style_name.lower()
    custom_prefixes = ["alchemy", "annexure", "schedule"]
    
    return any(style_lower.startswith(prefix) for prefix in custom_prefixes)

def _extract_template_styles(template_bytes: bytes, working_dir: Path) -> Dict[str, Any]:
    """
    Extract ONLY custom paragraph styles from template.
    Custom styles must be prefixed with: Alchemy, Annexure, or Schedule.
    """
    logging.info(f"🎨 Extracting custom template styles...")
    
    doc = Document(io.BytesIO(template_bytes))
    
    style_info = {}
    total_styles_found = 0
    custom_style_count = 0
    
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        
        total_styles_found += 1
        style_name = style.name
        
        # CRITICAL: Only process custom styles
        if not _is_custom_style(style_name):
            continue
        
        try:
            font = style.font
            
            properties = {
                "name": style_name,
                "is_custom": True,  # All extracted styles are custom
                "font_name": font.name if font.name else "Default",
                "font_size": float(font.size.pt) if font.size else None,
                "bold": font.bold if font.bold is not None else False,
                "italic": font.italic if font.italic is not None else False,
                "underline": font.underline if font.underline is not None else False,
                "color": "Default"
            }
            
            custom_style_count += 1
            
            # Get color
            if font.color and font.color.rgb:
                try:
                    rgb = font.color.rgb
                    properties["color"] = f"RGB({rgb[0]}, {rgb[1]}, {rgb[2]})"
                except:
                    pass
            
            # Get alignment
            if hasattr(style, 'paragraph_format') and style.paragraph_format:
                pf = style.paragraph_format
                if pf.alignment:
                    alignment_map = {
                        WD_ALIGN_PARAGRAPH.LEFT: "left",
                        WD_ALIGN_PARAGRAPH.CENTER: "center",
                        WD_ALIGN_PARAGRAPH.RIGHT: "right",
                        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
                    }
                    properties["alignment"] = alignment_map.get(pf.alignment, "left")
                else:
                    properties["alignment"] = "left"
            else:
                properties["alignment"] = "left"
            
            # Generate description
            desc_parts = []
            
            desc_parts.append("🎨 CUSTOM")
            
            if properties["font_size"]:
                desc_parts.append(f"{properties['font_size']}pt")
            
            if properties["bold"]:
                desc_parts.append("bold")
            
            if properties["italic"]:
                desc_parts.append("italic")
            
            if properties["alignment"] != "left":
                desc_parts.append(properties["alignment"])
            
            properties["description"] = ", ".join(desc_parts) if desc_parts else "custom style"
            
            style_info[style_name] = properties
            
        except Exception as e:
            logging.warning(f"  ⚠️  Could not extract style '{style_name}': {e}")
            continue
    
    logging.info(f"  ✅ Found {total_styles_found} total styles, extracted {custom_style_count} CUSTOM styles")
    
    if custom_style_count == 0:
        logging.error(f"  ❌ NO CUSTOM STYLES FOUND! Template must have styles prefixed with: Alchemy, Annexure, or Schedule")
    
    # Save styles for debugging
    output_file = working_dir / "template_styles.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump({
            "custom_styles": style_info,
            "total_styles_in_template": total_styles_found,
            "custom_styles_extracted": custom_style_count
        }, f, indent=2)
    
    return style_info

# ============================================================================
# NEW: STYLE MAPPING
# ============================================================================
def _find_custom_fallback_style(content_type: str, style_info: Dict[str, Any]) -> str:
    """
    Find the best custom style for a content type.
    Only searches within custom styles (Alchemy/Annexure/Schedule prefixed).
    """
    # Define preferred custom style names for each content type
    custom_preferences = {
        "heading_1": [
            "Alchemy Heading 1", "Alchemy Title", "Alchemy H1",
            "Annexure Heading 1", "Schedule Heading 1"
        ],
        "heading_2": [
            "Alchemy Heading 2", "Alchemy Subtitle", "Alchemy H2",
            "Annexure Heading 2", "Schedule Heading 2"
        ],
        "heading_3": [
            "Alchemy Heading 3", "Alchemy H3",
            "Annexure Heading 3", "Schedule Heading 3"
        ],
        "paragraph": [
            "Alchemy Body Text", "Alchemy Normal", "Alchemy Paragraph", "Alchemy Text",
            "Annexure Text", "Schedule Text"
        ],
        "list_item": [
            "Alchemy List", "Alchemy List Paragraph", "Alchemy Bullet",
            "Annexure List", "Schedule List"
        ],
        "citation_block": [
            "Alchemy Definitions", "Alchemy Citation", "Alchemy Reference",
            "Annexure Definitions", "Schedule Definitions"
        ]
    }
    
    # Try to find preferred custom styles first (case-insensitive)
    preferences = custom_preferences.get(content_type, [])
    for pref in preferences:
        # Case-insensitive match
        for style_name in style_info.keys():
            if style_name.lower() == pref.lower():
                logging.info(f"     🎨 Found custom preference: {style_name} for {content_type}")
                return style_name
    
    # Look for any custom style with relevant keywords
    keywords = {
        "heading_1": ["heading 1", "title", "h1", "header 1"],
        "heading_2": ["heading 2", "subtitle", "h2", "header 2"],
        "heading_3": ["heading 3", "h3", "header 3"],
        "paragraph": ["body", "normal", "text", "paragraph"],
        "list_item": ["list", "bullet", "item"],
        "citation_block": ["definition", "citation", "reference", "footnote"]
    }
    
    content_keywords = keywords.get(content_type, [])
    for style_name in style_info.keys():
        style_lower = style_name.lower()
        if any(keyword in style_lower for keyword in content_keywords):
            logging.info(f"     🎨 Found custom keyword match: {style_name} for {content_type}")
            return style_name
    
    # If no specific match, use first available custom style
    if style_info:
        first_style = list(style_info.keys())[0]
        logging.warning(f"     ⚠️  No specific match for {content_type}, using first custom style: {first_style}")
        return first_style
    
    # Absolute last resort: return None (will be handled by caller)
    logging.error(f"     ❌ No custom styles available for {content_type}!")
    return None

def _map_content_types_to_styles(style_info: Dict[str, Any], working_dir: Path) -> Dict[str, str]:
    """
    Use LLM to map content types to CUSTOM template styles only.
    Only uses styles prefixed with: Alchemy, Annexure, or Schedule.
    """
    logging.info(f"🎯 Mapping content types to CUSTOM template styles...")
    
    if not style_info:
        logging.error("❌ No custom styles available! Cannot create mapping.")
        return _create_fallback_style_mapping(style_info)
    
    # Build custom styles list (all styles are custom at this point)
    custom_styles = []
    for style_name, props in style_info.items():
        desc = f"• **{style_name}**: {props['description']}"
        if props.get('font_size'):
            desc += f" ({props['font_size']}pt"
            if props.get('bold'):
                desc += ", bold"
            if props.get('italic'):
                desc += ", italic"
            desc += ")"
        custom_styles.append(desc)
    
    styles_text = "**🎨 AVAILABLE CUSTOM STYLES:**\n" + "\n".join(custom_styles)
    
    custom_count = len(custom_styles)
    
    system_prompt = f"""You are a legal document formatting expert. Map legal opinion content types to custom Word template styles.

**AVAILABLE CUSTOM STYLES IN TEMPLATE:**
{styles_text}

**CRITICAL RULES:**
1. 🎨 You can ONLY use the custom styles listed above
2. These are firm-branded styles (Alchemy/Annexure/Schedule prefixed)
3. ALL {custom_count} styles are custom - choose the most appropriate for each content type
4. Match styles based on semantic meaning and visual hierarchy

**LEGAL OPINION CONTENT TYPES TO MAP:**
- heading_1: Main section titles (EXECUTIVE SUMMARY, OPINION, REFERENCES)
- heading_2: Subsection titles (Issue 1, Standard of Review, Recommendations)
- heading_3: Sub-subsection titles (4.1 Conclusion, Counter-arguments)
- paragraph: Regular legal opinion body text
- list_item: Bulleted/numbered points in legal arguments
- citation_block: Reference lists and citations

**MAPPING STRATEGY:**
1. Look for styles with "Heading 1", "Heading 2", etc. in their names
2. For paragraphs, look for "Body Text", "Normal", "Text", or "Paragraph"
3. For lists, look for "List", "Bullet", or "Item"
4. For citations, look for "Definitions", "Citation", "Reference"
5. Consider font size and bold styling when matching hierarchy
6. If no perfect semantic match exists, choose based on visual hierarchy (size/bold)

**OUTPUT FORMAT (JSON):**
{{
  "style_mapping": {{
    "heading_1": "Alchemy Heading 1",
    "heading_2": "Alchemy Heading 2",
    "heading_3": "Alchemy Heading 3",
    "paragraph": "Alchemy Body Text",
    "list_item": "Alchemy List",
    "citation_block": "Alchemy Definitions"
  }},
  "reasoning": {{
    "heading_1": "Alchemy Heading 1 is appropriate for main sections - largest heading style",
    "heading_2": "Alchemy Heading 2 for subsections - medium heading",
    ...
  }}
}}

**IMPORTANT:** 
- All style names in mapping MUST be from the custom styles list above
- Use your best judgment to match content types to available custom styles
- Respond with ONLY valid JSON"""

    try:
        response = call_llm_with(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Create the content type to style mapping using ONLY the {custom_count} custom styles available in the template."}
            ],
            temperature=0,
            max_tokens=2000,
            model_deployment_env_var='OPINION_MODEL_DEPLOYMENT',
            model_version_env_var='OPINION_MODEL_VERSION'
        )
        
        result_text = response.strip()
        
        # Try to extract JSON from response
        start = result_text.find('{')
        end = result_text.rfind('}') + 1
        if start != -1 and end > start:
            json_str = result_text[start:end]
            result = json.loads(json_str)
        else:
            result = json.loads(result_text)
        
        mapping = result.get("style_mapping", {})
        reasoning = result.get("reasoning", {})
        
        # Validate all mapped styles exist in custom styles
        available_custom_styles = set(style_info.keys())
        for content_type, style_name in mapping.items():
            if style_name not in available_custom_styles:
                logging.warning(f"  ⚠️  Mapped style '{style_name}' not found in custom styles, using fallback")
                fallback_style = _find_custom_fallback_style(content_type, style_info)
                if fallback_style:
                    mapping[content_type] = fallback_style
                else:
                    logging.error(f"  ❌ No fallback available for {content_type}!")
        
        # Log the mapping
        logging.info(f"  ✅ Created CUSTOM style mapping:")
        for content_type, style_name in mapping.items():
            reason = reasoning.get(content_type, "")
            logging.info(f"     🎨 {content_type} → {style_name} ({reason})")
        
        # Save mapping for debugging
        output_file = working_dir / "style_mapping.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                "mapping": mapping, 
                "reasoning": reasoning,
                "custom_styles_available": custom_count,
                "all_custom": True
            }, f, indent=2)
        
        return mapping
    
    except Exception as e:
        logging.error(f"  ❌ Style mapping failed: {e}")
        return _create_fallback_style_mapping(style_info)

def _create_fallback_style_mapping(style_info: Dict[str, Any]) -> Dict[str, str]:
    """
    Create fallback style mapping using ONLY custom styles.
    """
    logging.info(f"  🔧 Creating fallback CUSTOM style mapping...")
    
    if not style_info:
        logging.error("  ❌ No custom styles available! Cannot create fallback mapping.")
        # Return empty mapping - this will cause issues but alerts the user
        return {
            "heading_1": None,
            "heading_2": None,
            "heading_3": None,
            "paragraph": None,
            "list_item": None,
            "citation_block": None
        }
    
    content_types = ["heading_1", "heading_2", "heading_3", "paragraph", "list_item", "citation_block"]
    mapping = {}
    
    for content_type in content_types:
        fallback = _find_custom_fallback_style(content_type, style_info)
        if fallback:
            mapping[content_type] = fallback
        else:
            # Use first available custom style as absolute fallback
            first_custom = list(style_info.keys())[0]
            logging.warning(f"  ⚠️  No suitable custom style for {content_type}, using: {first_custom}")
            mapping[content_type] = first_custom
    
    # Log what we selected
    logging.info(f"  🎨 Fallback mapping uses ONLY custom styles:")
    for content_type, style_name in mapping.items():
        logging.info(f"     {content_type} → {style_name}")
    
    return mapping

# ============================================================================
# NEW: APPLY STYLES AND ASSEMBLE DOCUMENT
# ============================================================================
def _apply_styles_to_elements(elements: List[Dict[str, Any]], style_mapping: Dict[str, str]) -> List[Dict[str, Any]]:
    """
    Apply style assignments to all elements based on content type.
    """
    logging.info(f"🎨 Applying styles to {len(elements)} elements...")
    
    styled_elements = []
    
    for element in elements:
        content_type = element.get('type', 'paragraph')
        assigned_style = style_mapping.get(content_type, 'Normal')
        
        styled_element = {
            'content': element['content'],
            'type': content_type,
            'style': assigned_style
        }
        
        styled_elements.append(styled_element)
    
    # Log style distribution
    style_counts = {}
    for elem in styled_elements:
        style = elem['style']
        style_counts[style] = style_counts.get(style, 0) + 1
    
    logging.info(f"  ✅ Style distribution:")
    for style, count in sorted(style_counts.items(), key=lambda x: x[1], reverse=True):
        logging.info(f"     {style}: {count} elements")
    
    return styled_elements

def _preserve_template_sections(doc: Document) -> Dict[str, Any]:
    """
    Extract and preserve header, footer, and first-page content from template.
    Returns dict with preserved sections.
    """
    logging.info("📋 Preserving template header/footer...")
    
    preserved = {
        "has_header": False,
        "has_footer": False,
        "header_content": [],
        "footer_content": []
    }
    
    try:
        # Check for headers
        for section in doc.sections:
            if section.header:
                preserved["has_header"] = True
                for para in section.header.paragraphs:
                    if para.text.strip():
                        preserved["header_content"].append(para.text)
            
            if section.footer:
                preserved["has_footer"] = True
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        preserved["footer_content"].append(para.text)
            
            break  # Only check first section
        
        logging.info(f"  ✅ Header: {preserved['has_header']}, Footer: {preserved['has_footer']}")
        
    except Exception as e:
        logging.warning(f"  ⚠️  Could not preserve template sections: {e}")
    
    return preserved


# ============================================================================
# NEW: TEMPLATE STRUCTURE ANALYSIS
# ============================================================================
def _analyze_template_structure(template_bytes: bytes, working_dir: Path) -> Dict[str, Any]:
    """
    Analyze template to extract:
    1. Header section (To/From/Date/Re)
    2. Style references (which custom styles are mentioned)
    3. Footer sections (boilerplate, disclaimers, signatures)
    """
    logging.info(f"🔍 Analyzing template structure...")
    
    doc = Document(io.BytesIO(template_bytes))
    
    # Extract all body text
    body_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    
    logging.info(f"  📝 Template body text: {len(body_text):,} chars")
    
    system_prompt = """You are a legal document template analyzer.

**YOUR TASK:**
Analyze this Word template and identify its structural components.

**TEMPLATE STRUCTURE TO IDENTIFY:**

1. **HEADER SECTION**: The opening fields for client information
   - Look for: To:, From:, Date:, Re: fields
   - Usually at the very top of the document
   - May have placeholder text like [●], [Insert Subject], etc.
   - Extract the exact format and layout

2. **STYLE REFERENCE SECTION**: Lines that show available custom styles
   - These are lines that list style names (often with "Alchemy", "Annexure", "Schedule" prefixes)
   - Example: "Alchemy CLAUSE HEADING", "Alchemy Heading 2", "Indent 1"
   - Usually after header, before main content
   - These show which styles the template is designed to use
   - May include indentation examples

3. **BOILERPLATE/FOOTER SECTIONS**: Standard text that should appear in every opinion
   - Look for sections titled "General", "Limitations", "Disclaimers"  
   - Standard firm language about scope, liability, confidentiality
   - Usually at the bottom of the template
   - Includes signature blocks (firm name, location like "Alchemy\\nJohannesburg")
   - May have notes like "[Note: Delete below if not relevant...]"

**OUTPUT FORMAT (JSON):**
{
  "header_section": {
    "present": true,
    "raw_text": "To:\t[●] [Client entity name]\\nFrom:\tAlchemy Law\\nDate:\t[●]\\nRe:\t[Insert Subject]",
    "has_to_field": true,
    "has_from_field": true,
    "has_date_field": true,
    "has_re_field": true
  },
  "style_references": {
    "present": true,
    "raw_text": "Full text of the style reference section",
    "mentioned_styles": [
      "Alchemy CLAUSE HEADING",
      "Alchemy Heading 2",
      "Alchemy Heading 3",
      "Alchemy Heading 4",
      "Alchemy Heading 5",
      "Alchemy Heading 6"
    ],
    "notes": "These are the custom styles explicitly listed in the template"
  },
  "footer_sections": [
    {
      "title": "General",
      "content": "The opinion and advice set out herein are given solely...",
      "is_required": true
    },
    {
      "title": "Signature",
      "content": "Alchemy\\nJohannesburg",
      "is_required": true
    }
  ],
  "template_notes": "Any observations about the template structure"
}

**CRITICAL RULES:**
- Extract exact text from each section, don't paraphrase
- Preserve all placeholder markers like [●], [Insert...]
- List ALL custom style names found in the style reference section
- Include ALL footer sections in order
- Mark footer sections as required if they should appear in every opinion
- Return ONLY valid JSON"""

    try:
        response = call_llm_with(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Analyze this legal opinion template structure:\n\n{body_text}"}
            ],
            temperature=0,
            max_tokens=4000,
            model_deployment_env_var='OPINION_MODEL_DEPLOYMENT',
            model_version_env_var='OPINION_MODEL_VERSION'
        )
        
        # Parse response
        start = response.find('{')
        end = response.rfind('}') + 1
        if start != -1 and end > start:
            json_str = response[start:end]
            result = json.loads(json_str)
        else:
            result = json.loads(response)
        
        # Log findings
        header = result.get('header_section', {})
        style_refs = result.get('style_references', {})
        footers = result.get('footer_sections', [])
        
        logging.info(f"  ✅ Template structure analyzed:")
        logging.info(f"     📋 Header present: {header.get('present', False)}")
        if header.get('present'):
            logging.info(f"        • To field: {header.get('has_to_field', False)}")
            logging.info(f"        • From field: {header.get('has_from_field', False)}")
            logging.info(f"        • Date field: {header.get('has_date_field', False)}")
            logging.info(f"        • Re field: {header.get('has_re_field', False)}")
        
        mentioned_styles = style_refs.get('mentioned_styles', [])
        logging.info(f"     🎨 Style references found: {len(mentioned_styles)}")
        for style in mentioned_styles[:10]:  # Log first 10
            logging.info(f"        • {style}")
        if len(mentioned_styles) > 10:
            logging.info(f"        • ... and {len(mentioned_styles) - 10} more")
        
        logging.info(f"     📄 Footer sections: {len(footers)}")
        for footer in footers:
            required = footer.get('is_required', False)
            marker = "✅" if required else "ℹ️"
            logging.info(f"        {marker} {footer.get('title', 'Untitled')}")
        
        # Save for debugging
        output_file = working_dir / "template_structure.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
        
        return result
        
    except Exception as e:
        logging.error(f"  ❌ Template structure analysis failed: {e}")
        return _create_fallback_structure_analysis()

def _create_fallback_structure_analysis() -> Dict[str, Any]:
    """Create fallback structure if analysis fails."""
    logging.warning("  🔧 Creating fallback template structure")
    
    return {
        "header_section": {
            "present": False
        },
        "style_references": {
            "present": False,
            "mentioned_styles": []
        },
        "footer_sections": [],
        "template_notes": "Fallback structure due to analysis failure"
    }

# ============================================================================
# NEW: FILTER STYLES BY TEMPLATE REFERENCES
# ============================================================================
def _filter_styles_by_template_references(
    all_custom_styles: Dict[str, Any],
    template_structure: Dict[str, Any],
    working_dir: Path
) -> Dict[str, Any]:
    """
    Filter custom styles to only include those mentioned in the template.
    This ensures we only use styles the template is designed for.
    """
    logging.info(f"🔍 Filtering custom styles based on template references...")
    
    style_refs = template_structure.get('style_references', {})
    mentioned_styles = style_refs.get('mentioned_styles', [])
    
    if not mentioned_styles:
        logging.warning("  ⚠️  No style references found in template, using ALL custom styles")
        return all_custom_styles
    
    # Normalize mentioned style names for matching
    mentioned_normalized = {}
    for style in mentioned_styles:
        normalized = style.lower().strip().replace(' ', '').replace('-', '').replace('_', '')
        mentioned_normalized[normalized] = style
    
    filtered_styles = {}
    skipped_styles = []
    
    for style_name, style_props in all_custom_styles.items():
        # Normalize for comparison
        style_normalized = style_name.lower().strip().replace(' ', '').replace('-', '').replace('_', '')
        
        # Check for match
        is_matched = False
        for norm_key, original_mentioned in mentioned_normalized.items():
            if norm_key in style_normalized or style_normalized in norm_key:
                filtered_styles[style_name] = style_props
                logging.info(f"     ✅ Including: {style_name} (matches '{original_mentioned}' from template)")
                is_matched = True
                break
        
        if not is_matched:
            skipped_styles.append(style_name)
            logging.info(f"     ⏭️  Skipping: {style_name} (not referenced in template)")
    
    logging.info(f"  🎨 Filtered to {len(filtered_styles)}/{len(all_custom_styles)} styles")
    
    if not filtered_styles:
        logging.warning("  ⚠️  Filtering resulted in zero styles! Reverting to all custom styles")
        filtered_styles = all_custom_styles
    
    # Save filtered results
    output_file = working_dir / "filtered_custom_styles.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump({
            "mentioned_in_template": mentioned_styles,
            "filtered_styles": list(filtered_styles.keys()),
            "skipped_styles": skipped_styles,
            "total_custom_styles": len(all_custom_styles),
            "styles_after_filtering": len(filtered_styles)
        }, f, indent=2)
    
    return filtered_styles

# ============================================================================
# NEW: POPULATE HEADER SECTION
# ============================================================================
def _populate_header_section(
    template_structure: Dict[str, Any],
    client_name: str,
    date_str: str,
    subject: str
) -> List[str]:
    """
    Generate header lines with populated fields.
    Returns list of lines to add at document start.
    """
    header = template_structure.get('header_section', {})
    
    if not header.get('present'):
        logging.info("  ℹ️  No header section in template")
        return []
    
    logging.info(f"  📋 Populating header section...")
    
    # Start with the raw template text
    header_text = header.get('raw_text', '')
    
    # Replace placeholders with actual values
    populated = header_text
    
    # Replace To: field - handle various placeholder formats
    if header.get('has_to_field'):
        populated = re.sub(
            r'(To:\s*)\[●\]\s*\[Client entity name\]',
            f'To:\t{client_name}',
            populated,
            flags=re.IGNORECASE
        )
        populated = re.sub(
            r'(To:\s*)\[.*?\]',
            f'To:\t{client_name}',
            populated,
            flags=re.IGNORECASE
        )
    
    # Date: field
    if header.get('has_date_field'):
        populated = re.sub(
            r'(Date:\s*)\[●\]',
            f'Date:\t{date_str}',
            populated,
            flags=re.IGNORECASE
        )
        populated = re.sub(
            r'(Date:\s*)\[.*?\]',
            f'Date:\t{date_str}',
            populated,
            flags=re.IGNORECASE
        )
    
    # Re: field
    if header.get('has_re_field'):
        populated = re.sub(
            r'(Re:\s*)\[Insert Subject\]',
            f'Re:\t{subject}',
            populated,
            flags=re.IGNORECASE
        )
        populated = re.sub(
            r'(Re:\s*)\[.*?\]',
            f'Re:\t{subject}',
            populated,
            flags=re.IGNORECASE
        )
    
    # Split into lines
    header_lines = [line.strip() for line in populated.split('\n') if line.strip()]
    
    logging.info(f"  ✅ Generated {len(header_lines)} header lines")
    
    return header_lines

# ============================================================================
# NEW: EXTRACT FOOTER SECTIONS
# ============================================================================
def _extract_footer_content(template_structure: Dict[str, Any]) -> List[Dict[str, str]]:
    """
    Extract footer sections from template structure.
    Returns list of footer sections to append at document end.
    """
    footer_sections = template_structure.get('footer_sections', [])
    
    if not footer_sections:
        logging.info("  ℹ️  No footer sections in template")
        return []
    
    logging.info(f"  📄 Extracting {len(footer_sections)} footer section(s)...")
    
    processed_footers = []
    
    for footer in footer_sections:
        title = footer.get('title', '')
        content = footer.get('content', '')
        is_required = footer.get('is_required', True)
        
        if is_required and content.strip():
            processed_footers.append({
                "title": title,
                "content": content.strip()
            })
            logging.info(f"     ✅ Including footer: {title}")
        else:
            logging.info(f"     ⏭️  Skipping optional footer: {title}")
    
    return processed_footers

def _assemble_document_with_template_structure(
    template_bytes: bytes,
    template_structure: Dict[str, Any],
    styled_elements: List[Dict[str, Any]],
    style_info: Dict[str, Any],
    client_name: str,
    date_str: str,
    subject: str,
    workspace_dir: Path,
    filename_base: str
) -> bytes:
    """
    Assemble document with full template structure:
    1. Header (To/From/Date/Re)
    2. Opinion content (styled elements)
    3. Footer sections (boilerplate, disclaimers, signature)
    """
    logging.info(f"📄 Assembling document with complete template structure...")
    
    # Load template
    doc = Document(io.BytesIO(template_bytes))
    
    # Preserve header/footer from document properties
    preserved = _preserve_template_sections(doc)
    
    # Clear existing body content
    logging.info(f"  🧹 Clearing template body content...")
    for para in reversed(doc.paragraphs):
        p_element = para._element
        p_element.getparent().remove(p_element)
    
    logging.info(f"  ✅ Template body cleared")
    
    available_styles = set(style_info.keys())
    
    # ========================================================================
    # SECTION 1: Add Header
    # ========================================================================
    header_lines = _populate_header_section(template_structure, client_name, date_str, subject)
    
    if header_lines:
        logging.info(f"  📋 Adding header section ({len(header_lines)} lines)...")
        for line in header_lines:
            para = doc.add_paragraph(line)
            # Use Normal style for header lines
            if 'Normal' in available_styles:
                try:
                    para.style = 'Normal'
                except:
                    pass
        
        # Add spacing after header
        doc.add_paragraph("")  # Blank line
        logging.info(f"     ✅ Header added")
    
    # ========================================================================
    # SECTION 2: Add Opinion Content (styled elements)
    # ========================================================================
    logging.info(f"  📝 Adding opinion content ({len(styled_elements)} elements)...")
    
    for idx, element in enumerate(styled_elements):
        content = element['content']
        style_name = element['style']
        
        para = doc.add_paragraph(content)
        
        # Apply custom style
        if style_name and style_name in available_styles:
            try:
                para.style = style_name
            except Exception as e:
                logging.warning(f"  ⚠️  Failed to apply style '{style_name}': {e}")
        
        # Progress logging
        if (idx + 1) % 100 == 0:
            logging.info(f"     Added {idx + 1}/{len(styled_elements)} elements...")
    
    logging.info(f"     ✅ Opinion content added")
    
    # ========================================================================
    # SECTION 3: Add Footer Sections
    # ========================================================================
    footer_sections = _extract_footer_content(template_structure)
    
    if footer_sections:
        logging.info(f"  📄 Adding footer sections ({len(footer_sections)} sections)...")
        
        # Add spacing before footer
        doc.add_paragraph("")  # Blank line
        
        for footer in footer_sections:
            title = footer.get('title', '')
            content = footer.get('content', '')
            
            # Add title as heading if present
            if title:
                title_para = doc.add_paragraph(title)
                # Try to use a heading style for footer title
                if 'Alchemy Heading 3' in available_styles:
                    try:
                        title_para.style = 'Alchemy Heading 3'
                    except:
                        pass
            
            # Add content
            content_para = doc.add_paragraph(content)
            if 'Normal' in available_styles:
                try:
                    content_para.style = 'Normal'
                except:
                    pass
            
            # Add spacing after each footer section
            doc.add_paragraph("")
            
            logging.info(f"     ✅ Added footer: {title}")
    
    # Save to bytes
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_bytes = output_buffer.getvalue()
    
    logging.info(f"  ✅ Document assembled with template structure ({len(output_bytes):,} bytes)")
    
    return output_bytes

# ============================================================================
# NEW: DOCUMENT CLEANUP
# ============================================================================
def _remove_numbering_from_text(text: str) -> Tuple[str, bool]:
    """
    Remove duplicate numbering patterns from text that may conflict with template.
    Returns (cleaned_text, was_modified)
    """
    original = text
    
    # Pattern 1: Remove leading numbers with dots (1.1, 2.3.4, etc.) - but preserve [1] citations
    text = re.sub(r'^\s*\d+(\.\d+)+\.?\s+', '', text)
    
    # Pattern 2: Remove leading single numbers followed by period
    text = re.sub(r'^\s*\d+\.\s+(?!\[)', '', text)
    
    # Pattern 3: Remove leading letters in parentheses
    text = re.sub(r'^\s*\([a-zA-Z]\)\s+', '', text)
    
    # Pattern 4: Remove "Chapter/Section X" at start
    text = re.sub(r'^\s*(Chapter|Section|Part)\s+\d+[\.:]\s*', '', text, flags=re.IGNORECASE)
    
    return text.strip(), text != original

def _review_and_cleanup_document(doc_bytes: bytes, working_dir: Path) -> bytes:
    """
    Review and clean up the assembled document:
    1. Remove duplicate numbering from headings
    2. Clean up formatting issues
    
    Returns bytes of cleaned document.
    """
    logging.info(f"🔍 Reviewing and cleaning up document...")
    
    # Load document
    doc = Document(io.BytesIO(doc_bytes))
    
    # Statistics
    numbering_removed_count = 0
    
    # Process each paragraph
    logging.info(f"  🔄 Processing {len(doc.paragraphs)} paragraphs...")
    
    for idx, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        
        # Remove duplicate numbering from text (but preserve citations)
        original_text = para.text
        cleaned_text, was_modified = _remove_numbering_from_text(original_text)
        
        if was_modified:
            # Replace paragraph text while preserving formatting
            para.text = cleaned_text
            numbering_removed_count += 1
            
            if idx < 10:  # Log first 10 for debugging
                logging.info(f"     ✂️  Cleaned: {original_text[:60]} → {cleaned_text[:60]}")
    
    # Save cleaned document
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    cleaned_bytes = output_buffer.getvalue()
    
    logging.info(f"  ✅ Document cleanup complete:")
    logging.info(f"     • Removed numbering from {numbering_removed_count} paragraphs")
    logging.info(f"     • Final size: {len(cleaned_bytes):,} bytes")
    
    return cleaned_bytes

def _get_current_draft_text(opinion: dict) -> str:
    st = opinion.get("staging_draft")
    if st:
        d = st.get("draft")
        if isinstance(d, str):
            return d
        if isinstance(d, dict) and "draft" in d:
            return d["draft"]
    raise ValueError("No draft text found for this opinion.")

def format_opinion_for_docxtpl(opinion_text: str, client_name: str, subject: str) -> dict:
    system_prompt = f"""You are a legal document formatter specializing in converting legal opinions for DocxTemplate.

CRITICAL TEXT FORMATTING RULES:
1. Use single spaces between all words - no double spaces
2. Remove any extra whitespace or special characters
3. Ensure clean, consistent spacing throughout
4. No tabs or unusual spacing characters

Your task is to convert a legal opinion into a structured JSON format that will work perfectly with DocxTemplate RichText formatting.

The output must be a JSON object with this EXACT structure:
{{
  "body_sections": [
    {{
      "type": "heading",
      "text": "EXECUTIVE SUMMARY",
      "style": {{"bold": true, "size": 22}}
    }},
    {{
      "type": "paragraph", 
      "text": "Content text here with proper single spacing between words.",
      "style": {{"size": 16}}
    }}
  ]
}}

FORMATTING REQUIREMENTS:
1. Convert **bold** markdown to proper heading sections with bold styling
2. Remove excessive line breaks and normalize spacing
3. ENSURE SINGLE SPACES BETWEEN ALL WORDS
4. Convert section headers like "**CLIENT INFORMATION**" to proper headings
5. Create bullet points where appropriate (use "• " format)
6. Remove citation numbers like [1], [2] 
7. Make the text client-friendly (less legal jargon)
8. Clean up any irregular spacing or formatting

CLIENT CONTEXT:
- Client: {client_name}
- Subject: {subject}

Make sure all text has consistent, single-space formatting.
Return ONLY valid JSON with the exact structure shown above."""

    user_prompt = f"""Please convert this legal opinion into the structured JSON format for DocxTemplate:

{opinion_text}"""  # Limit text to avoid token limits

    try:
        response = call_llm_with(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.0,
            max_tokens=8000,
            model_deployment_env_var='OPINION_MODEL_DEPLOYMENT',
            model_version_env_var='OPINION_MODEL_VERSION'
        )
        
        # Parse JSON response
        try:
            # Try to extract JSON from response
            start = response.find('{')
            end = response.rfind('}') + 1
            if start != -1 and end > start:
                json_str = response[start:end]
                return json.loads(json_str)
            else:
                return json.loads(response)
        except json.JSONDecodeError:
            logging.error(f"Failed to parse LLM JSON response: {response[:500]}")
            # Return fallback structure
            return create_fallback_structure(opinion_text, client_name)
            
    except Exception as e:
        logging.error(f"Error in format_opinion_for_docxtpl: {e}")
        return create_fallback_structure(opinion_text, client_name)

def create_fallback_structure(opinion_text: str, client_name: str) -> dict:
    """Create a fallback structured format if LLM formatting fails"""
    
    # Basic cleanup
    cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'\1', opinion_text)  # Remove markdown bold
    cleaned_text = re.sub(r'\[(\d+)\]', '', cleaned_text)  # Remove citations
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)  # Fix line breaks
    
    return {
        "body_sections": [
            {
                "type": "heading",
                "text": "LEGAL OPINION",
                "style": {"bold": True, "size": 26}
            },
            {
                "type": "paragraph",
                "text": f"Dear {client_name},",
                "style": {"size": 22}
            },
            {
                "type": "paragraph", 
                "text": cleaned_text,
                "style": {"size": 22}
            }
        ]
    }

def convert_to_richtext_body(structured_data: dict) -> RichText:
    """Convert structured data to a RichText object with proper spacing and cleaned text"""
    
    rt = RichText()
    sections = structured_data.get("body_sections", [])
    
    for i, section in enumerate(sections):
        section_type = section.get("type", "paragraph")
        text = section.get("text", "")
        style = section.get("style", {})
        
        if not text.strip():
            continue
        
        # Clean the text to remove spacing issues
        cleaned_text = clean_text_for_word(text)
        
        if section_type == "heading":
            # Add paragraph break before headings (except first)
            if i > 0:
                rt.add('\a')  # New paragraph break
            
            rt.add(
                cleaned_text,
                bold=style.get("bold", True),
                size=style.get("size", 26)
            )
            rt.add('\a')  # Single paragraph break after heading
            
        elif section_type == "list_item":
            rt.add(
                cleaned_text,
                size=style.get("size", 22)
            )
            rt.add('\n')  # Just a line break for list items
            
        else:  # paragraph
            rt.add(
                cleaned_text,
                size=style.get("size", 22),
                bold=style.get("bold", False)
            )
            rt.add('\a')  # Paragraph break for regular paragraphs
    
    return rt

def clean_text_for_word(text: str) -> str:
    """Clean text to prevent Word spacing issues"""
    
    # Remove excessive whitespace
    cleaned = re.sub(r'\s+', ' ', text.strip())
    
    # Remove any hidden characters that could cause spacing issues
    cleaned = re.sub(r'[\u200B\u200C\u200D\uFEFF]', '', cleaned)  # Zero-width characters
    cleaned = re.sub(r'[\u00A0]', ' ', cleaned)  # Non-breaking spaces to regular spaces
    
    # Remove multiple spaces between words
    cleaned = re.sub(r' +', ' ', cleaned)
    
    # Remove leading/trailing spaces around punctuation
    cleaned = re.sub(r' +([.,;:!?])', r'\1', cleaned)
    cleaned = re.sub(r'([.,;:!?]) +', r'\1 ', cleaned)
    
    return cleaned

def render_docx_with_richtext(template_bytes: bytes, context: dict) -> bytes:
    """Render DocxTemplate with RichText formatting"""
    
    logging.info(f"Rendering DocxTemplate with RichText context keys: {list(context.keys())}")
    
    try:
        tpl = DocxTemplate(io.BytesIO(template_bytes))
        tpl.render(context)
        
        buf = io.BytesIO()
        tpl.save(buf)
        result_bytes = buf.getvalue()
        
        logging.info(f"DocxTemplate rendered successfully. Output size: {len(result_bytes)} bytes")
        return result_bytes
        
    except Exception as e:
        logging.error(f"DocxTemplate rendering failed: {e}")
        logging.error(f"Context was: {context}")
        raise

def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info("✅ OpinionCompile with Template Population triggered")
    
    if req.headers.get("function-key") != os.environ["FUNCTION_KEY"]:
        return func.HttpResponse("", status_code=401)
    
    try:
        body = req.get_json()
        email, err = auth_get_email(req)
        if err:
            return err

        # Extract parameters
        opinion_id = body.get("opinion_id")
        to_field = body.get("to", "")
        date_field = body.get("date", "")
        re_field = body.get("re", "")
        provided_draft = body.get("staging_draft_text")
        template_b64 = body.get("template_docx_b64")
        template_name = body.get("template_filename", "uploaded_template.docx")
        
        if not opinion_id:
            return func.HttpResponse("opinion_id is required", status_code=400)

        # Load opinion data
        user = get_user_info(email)
        payload = user.get("clean_payload") or json.loads(user.get("payload", "{}"))
        opinion = next((o for o in payload.get("opinions", []) if o.get("id") == opinion_id), None)
        
        if not opinion:
            return func.HttpResponse("Opinion not found", status_code=404)

        # Get draft text
        if isinstance(provided_draft, str) and provided_draft.strip():
            draft_text = provided_draft.strip()
            logging.info("📄 Using provided draft text")
        else:
            draft_text = _get_current_draft_text(opinion)
            logging.info("📄 Using staging draft text")

        logging.info(f"📄 Original lawyer draft length: {len(draft_text)}")

        # Add numbered citations to opinion
        logging.info("📚 Adding numbered citations to opinion...")
        draft_text_with_citations = add_citations_to_opinion(draft_text)
        logging.info(f"📄 Opinion with citations length: {len(draft_text_with_citations)}")

        # Prepare context data
        client_name = opinion.get("client_name", "") or to_field or "Client"
        subject = opinion.get("title", "") or re_field or "Legal Opinion"
        facts = opinion.get("facts", "")
        questions = opinion.get("questions", "")
        assumptions = opinion.get("assumptions", "")
        
        if not date_field:
            date_field = datetime.now(TZ_SA).strftime("%d %B %Y")

        # ====================================================================
        # DECISION POINT: Template-based or Simple RichText?
        # ====================================================================
        
        if template_b64:
            # ================================================================
            # TEMPLATE POPULATION PIPELINE
            # ================================================================
            logging.info("="*80)
            logging.info("🎨 TEMPLATE POPULATION PIPELINE")
            logging.info("="*80)
            
            import base64
            import tempfile
            
            start_time = datetime.now()
            
            # Decode template
            template_bytes = base64.b64decode(template_b64)
            logging.info(f"📎 Using uploaded template: {template_name} ({len(template_bytes):,} bytes)")
            
            # Create working directory
            with tempfile.TemporaryDirectory() as temp_dir:
                working_dir = Path(temp_dir)
                logging.info(f"📁 Working directory: {working_dir}")
                
                # Convert to client-friendly version
                logging.info("🎯 Converting to comprehensive client version...")
                client_structured_data = convert_lawyer_opinion_to_client_version(
                    draft_text_with_citations,
                    client_name,
                    subject,
                    facts,
                    questions,
                    assumptions
                )
                
                # Extract text from structured data for processing
                client_opinion_text = ""
                for section in client_structured_data.get('body_sections', []):
                    text = section.get('text', '')
                    section_type = section.get('type', 'paragraph')
                    
                    # Add markdown formatting hints for structure detection
                    if section_type == 'heading':
                        client_opinion_text += f"\n\n## {text}\n\n"
                    elif section_type == 'list_item':
                        client_opinion_text += f"- {text}\n"
                    else:
                        client_opinion_text += f"{text}\n\n"
                
                logging.info(f"📝 Client opinion text prepared: {len(client_opinion_text):,} chars")
                
                # Phase 1: Extract structure from opinion
                logging.info("\n" + "="*60)
                logging.info("📄 [Phase 1] Extracting Structure from Opinion")
                logging.info("="*60)
                
                phase1_start = datetime.now()
                structured_elements = _process_opinion_text(client_opinion_text, working_dir)
                phase1_time = (datetime.now() - phase1_start).total_seconds()
                logging.info(f"✅ Phase 1 complete in {phase1_time:.2f}s")
                
                # Phase 2: Analyze template structure
                logging.info("\n" + "="*60)
                logging.info("🔍 [Phase 2] Analyzing Template Structure")
                logging.info("="*60)
                
                phase2_start = datetime.now()
                template_structure = _analyze_template_structure(template_bytes, working_dir)
                phase2_time = (datetime.now() - phase2_start).total_seconds()
                logging.info(f"✅ Phase 2 complete in {phase2_time:.2f}s")
                
                # Phase 3: Extract ALL custom template styles
                logging.info("\n" + "="*60)
                logging.info("🎨 [Phase 3] Extracting ALL Custom Template Styles")
                logging.info("="*60)
                
                phase3_start = datetime.now()
                all_custom_styles = _extract_template_styles(template_bytes, working_dir)
                phase3_time = (datetime.now() - phase3_start).total_seconds()
                logging.info(f"✅ Phase 3 complete in {phase3_time:.2f}s")
                
                # Phase 3.5: Filter styles based on template references
                logging.info("\n" + "="*60)
                logging.info("🔍 [Phase 3.5] Filtering Styles by Template References")
                logging.info("="*60)
                
                phase3_5_start = datetime.now()
                style_info = _filter_styles_by_template_references(
                    all_custom_styles=all_custom_styles,
                    template_structure=template_structure,
                    working_dir=working_dir
                )
                phase3_5_time = (datetime.now() - phase3_5_start).total_seconds()
                logging.info(f"✅ Phase 3.5 complete in {phase3_5_time:.2f}s")
                
                # Phase 4: Map content types to filtered styles
                logging.info("\n" + "="*60)
                logging.info("🎯 [Phase 4] Mapping Content Types to Styles")
                logging.info("="*60)
                
                phase4_start = datetime.now()
                style_mapping = _map_content_types_to_styles(style_info, working_dir)
                phase4_time = (datetime.now() - phase4_start).total_seconds()
                logging.info(f"✅ Phase 4 complete in {phase4_time:.2f}s")
                
                # Phase 5: Apply styles to elements
                logging.info("\n" + "="*60)
                logging.info("🎨 [Phase 5] Applying Styles to Elements")
                logging.info("="*60)
                
                phase5_start = datetime.now()
                styled_elements = _apply_styles_to_elements(structured_elements, style_mapping)
                phase5_time = (datetime.now() - phase5_start).total_seconds()
                logging.info(f"✅ Phase 5 complete in {phase5_time:.2f}s")
                
                # Phase 6: Assemble document WITH TEMPLATE STRUCTURE
                logging.info("\n" + "="*60)
                logging.info("📄 [Phase 6] Assembling Document with Template Structure")
                logging.info("="*60)
                
                phase6_start = datetime.now()
                populated_bytes = _assemble_document_with_template_structure(
                    template_bytes=template_bytes,
                    template_structure=template_structure,
                    styled_elements=styled_elements,
                    style_info=style_info,
                    client_name=client_name,
                    date_str=date_field,
                    subject=subject,
                    workspace_dir=working_dir,
                    filename_base=Path(template_name).stem
                )
                phase6_time = (datetime.now() - phase6_start).total_seconds()
                logging.info(f"✅ Phase 6 complete in {phase6_time:.2f}s")
                
                # Phase 7: Cleanup
                logging.info("\n" + "="*60)
                logging.info("🔍 [Phase 7] Reviewing & Cleaning Up")
                logging.info("="*60)
                
                phase7_start = datetime.now()
                final_bytes = _review_and_cleanup_document(populated_bytes, working_dir)
                phase7_time = (datetime.now() - phase7_start).total_seconds()
                logging.info(f"✅ Phase 7 complete in {phase7_time:.2f}s")
            
            # Upload to blob storage
            date_str = datetime.now(TZ_SA).strftime("%Y%m%d")
            safe_client = re.sub(r'[^\w\-_]', '_', client_name or "Client")
            file_id = generate_identifier()
            filename = f"Opinion_{safe_client}_{date_str}_{file_id}.docx"
            
            container = os.environ.get("COMPILED_CONTAINER", "opiniondrafts")
            blob_path = f"{opinion_id}/{filename}"
            
            write_to_blob_storage(
                _conn_str(),
                container,
                blob_path,
                final_bytes,
                meta_data={
                    "original_file_name": filename,
                    "extension": "docx",
                    "opinion_id": opinion_id,
                    "client_name": client_name,
                    "version": "client_focused_template",
                    "template_used": template_name
                },
                overwrite=True
            )
            
            # Generate download URL
            url = get_blob_sas_url(_conn_str(), container, blob_path, expiry_minutes=60*24)
            
            total_time = (datetime.now() - start_time).total_seconds()
            
            logging.info(f"✅ Template-based opinion compiled successfully: {filename}")
            
            return func.HttpResponse(
                body=json.dumps({
                    "success": True,
                    "message": f"Opinion compiled using template '{template_name}' successfully",
                    "url": url,
                    "blob_path": f"{container}/{blob_path}",
                    "file_name": filename,
                    "version": "client_focused_template",
                    "template_name": template_name,
                    "processing_stats": {
                        "total_elements": len(styled_elements),
                        "styles_used": len(style_info),
                        "phase_times": {
                            "structure_extraction": f"{phase1_time:.1f}s",
                            "template_analysis": f"{phase2_time:.1f}s",
                            "style_extraction": f"{phase3_time:.1f}s",
                            "style_filtering": f"{phase3_5_time:.1f}s",
                            "style_mapping": f"{phase4_time:.1f}s",
                            "style_application": f"{phase5_time:.1f}s",
                            "document_assembly": f"{phase6_time:.1f}s",
                            "cleanup": f"{phase7_time:.1f}s",
                            "total": f"{total_time:.1f}s"
                        }
                    }
                }),
                mimetype="application/json",
                status_code=200
            )
        
        else:
            # ================================================================
            # ORIGINAL RICHTEXT PIPELINE (Default template)
            # ================================================================
            logging.info("🎯 Converting lawyer opinion to client version (RichText)...")
            
            client_structured_data = convert_lawyer_opinion_to_client_version(
                draft_text_with_citations,
                client_name,
                subject,
                facts,
                questions,
                assumptions
            )
            
            sections_count = len(client_structured_data.get('body_sections', []))
            logging.info(f"📋 Client version contains {sections_count} sections")
            
            # Convert to RichText (your existing function)
            richtext_body = convert_to_richtext_body(client_structured_data)
            
            logging.info("✨ Created client-focused RichText body")
            
            # Use default template
            t_container = os.environ["TEMPLATE_CONTAINER"]
            t_blob_name = os.environ["TEMPLATE_BLOB_NAME"]
            template_bytes = read_from_blob_storage(_conn_str(), t_container, t_blob_name)
            logging.info(f"📎 Using default template")
            
            # Create template context
            context = {
                "to": to_field or client_name,
                "from": "Alchemy Law",
                "date": date_field,
                "re": re_field or subject,
                "body": richtext_body
            }
            
            # Render template (your existing function)
            compiled_bytes = render_docx_with_richtext(template_bytes, context)
            
            # Upload to blob storage
            date_str = datetime.now(TZ_SA).strftime("%Y%m%d")
            safe_client = re.sub(r'[^\w\-_]', '_', client_name or "Client")
            file_id = generate_identifier()
            filename = f"Opinion_{safe_client}_{date_str}_{file_id}.docx"
            
            container = os.environ.get("COMPILED_CONTAINER", "opiniondrafts")
            blob_path = f"{opinion_id}/{filename}"
            
            write_to_blob_storage(
                _conn_str(),
                container,
                blob_path,
                compiled_bytes,
                meta_data={
                    "original_file_name": filename,
                    "extension": "docx",
                    "opinion_id": opinion_id,
                    "client_name": client_name,
                    "version": "client_focused"
                },
                overwrite=True
            )
            
            # Generate download URL
            url = get_blob_sas_url(_conn_str(), container, blob_path, expiry_minutes=60*24)
            
            logging.info(f"✅ Client-focused opinion compiled successfully: {filename}")
            
            return func.HttpResponse(
                body=json.dumps({
                    "success": True,
                    "message": "Comprehensive client opinion compiled successfully",
                    "url": url,
                    "blob_path": f"{container}/{blob_path}",
                    "file_name": filename,
                    "version": "client_focused",
                    "sections_count": sections_count
                }),
                mimetype="application/json",
                status_code=200
            )
    
    except Exception as e:
        logging.exception("❌ Error in OpinionCompile")
        return func.HttpResponse(
            body=json.dumps({
                "success": False,
                "error": str(e)
            }),
            mimetype="application/json",
            status_code=500
        )