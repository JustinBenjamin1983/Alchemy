"""
Document loader that extracts text from .docx files and classifies them.
"""
from pathlib import Path
from dataclasses import dataclass
from typing import List, Dict, Optional
import re

try:
    from docx import Document
except ImportError:
    Document = None


@dataclass
class LoadedDocument:
    """Represents a loaded document with metadata."""
    filename: str
    filepath: Path
    text: str
    word_count: int
    doc_type: str  # constitutional, contract, regulatory, financial, employment, governance

    @property
    def char_count(self) -> int:
        return len(self.text)

    def to_dict(self) -> Dict:
        """Convert to dictionary for processing."""
        return {
            "filename": self.filename,
            "filepath": str(self.filepath),
            "text": self.text,
            "word_count": self.word_count,
            "doc_type": self.doc_type,
            "char_count": self.char_count
        }


def classify_document(filename: str, text: str) -> str:
    """
    Classify document type based on filename and content.

    Types:
    - constitutional: MOI, Shareholders Agreement
    - governance: Board Resolutions, Minutes
    - contract: Commercial agreements
    - regulatory: Mining rights, licenses, permits
    - financial: AFS, financial statements
    - employment: Employment contracts
    """
    filename_lower = filename.lower()
    text_lower = text.lower()[:2000]  # Check first 2000 chars

    # Constitutional documents
    if "moi" in filename_lower or "memorandum of incorporation" in text_lower:
        return "constitutional"
    if "shareholder" in filename_lower and "agreement" in filename_lower:
        return "constitutional"

    # Governance
    if "board" in filename_lower and "resolution" in filename_lower:
        return "governance"
    if "minutes" in filename_lower:
        return "governance"

    # Employment
    if "employment" in filename_lower or "contract" in filename_lower:
        if "employment" in text_lower or "employee" in text_lower:
            return "employment"

    # Financial
    if "afs" in filename_lower or "financial statement" in filename_lower:
        return "financial"
    if "annual financial" in text_lower:
        return "financial"

    # Regulatory
    if "mining" in filename_lower and ("right" in filename_lower or "10158" in filename_lower):
        return "regulatory"
    if "environmental" in filename_lower:
        return "regulatory"
    if "license" in filename_lower or "permit" in filename_lower:
        return "regulatory"

    # Contracts (default for business documents)
    if "loan" in filename_lower or "facility" in filename_lower or "bank" in filename_lower:
        return "contract"
    if "lease" in filename_lower:
        return "contract"
    if "supply" in filename_lower or "agreement" in filename_lower:
        return "contract"
    if "eskom" in filename_lower:
        return "contract"

    return "other"


def extract_text_from_docx(filepath: Path) -> str:
    """Extract text from a .docx file."""
    if Document is None:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    doc = Document(filepath)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n\n".join(paragraphs)


def load_documents(doc_dir: Path) -> List[LoadedDocument]:
    """
    Load all .docx files from directory.

    Args:
        doc_dir: Path to directory containing .docx files

    Returns:
        List of LoadedDocument objects
    """
    if isinstance(doc_dir, str):
        doc_dir = Path(doc_dir)

    if not doc_dir.exists():
        raise FileNotFoundError(f"Document directory not found: {doc_dir}")

    documents = []

    for filepath in sorted(doc_dir.glob("*.docx")):
        # Skip temp files
        if filepath.name.startswith("~$"):
            continue

        try:
            text = extract_text_from_docx(filepath)

            loaded = LoadedDocument(
                filename=filepath.name,
                filepath=filepath,
                text=text,
                word_count=len(text.split()),
                doc_type=classify_document(filepath.name, text)
            )
            documents.append(loaded)

        except Exception as e:
            print(f"Warning: Error loading {filepath.name}: {e}")

    return documents


def get_reference_documents(documents: List[LoadedDocument]) -> List[LoadedDocument]:
    """
    Get constitutional/governance docs that should always be included in context.

    These are the "source of truth" documents that other documents should be
    validated against (MOI, Shareholders Agreement, Board Resolutions).
    """
    return [d for d in documents if d.doc_type in ("constitutional", "governance")]


def build_document_context(
    documents: List[LoadedDocument],
    include_types: Optional[List[str]] = None,
    max_chars: Optional[int] = None
) -> str:
    """
    Build a combined context string from multiple documents.

    Args:
        documents: List of documents to include
        include_types: If provided, only include these doc_types
        max_chars: Maximum total characters (truncates proportionally)

    Returns:
        Formatted string with all document content
    """
    filtered = documents
    if include_types:
        filtered = [d for d in documents if d.doc_type in include_types]

    sections = []
    for doc in filtered:
        sections.append(f"""
{'='*60}
DOCUMENT: {doc.filename}
TYPE: {doc.doc_type.upper()}
WORDS: {doc.word_count:,}
{'='*60}

{doc.text}
""")

    combined = "\n\n".join(sections)

    if max_chars and len(combined) > max_chars:
        # Truncate proportionally
        combined = combined[:max_chars] + "\n\n[... TRUNCATED ...]"

    return combined
