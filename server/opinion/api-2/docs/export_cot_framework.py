"""
Export Mining DD CoT Framework to Word Document
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def create_cot_framework_doc():
    """Convert the markdown framework to a professional Word document."""

    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, "Mining_DD_CoT_Framework.md")

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)

    # Heading styles
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)

    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)

    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(51, 102, 153)

    # Process content
    lines = content.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block and 'json' in line:
                # Add a label for JSON blocks
                p = doc.add_paragraph()
                run = p.add_run('JSON Output:')
                run.bold = True
                run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.3)
            # Light gray background simulation via spacing
            i += 1
            continue

        # Tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_rows = []

            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and all(set(c) <= set('-: ') for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            if table_rows:
                render_table(doc, table_rows)
            in_table = False
            table_rows = []
            continue

        # Main title
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_paragraph(line[2:], style='CustomTitle')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Subtitle
        if line.startswith('## ') and 'Chain-of-Thought' in line:
            p = doc.add_paragraph(line[3:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(102, 102, 102)
            i += 1
            continue

        # Section headers
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue

        if line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            i += 1
            continue

        # Horizontal rules
        if line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('─' * 80)
            p.runs[0].font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue

        # Bold text (standalone)
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            text = line[2:-2]
            run = p.add_run(text)
            run.bold = True
            i += 1
            continue

        # Bold label lines
        if line.startswith('**') and ':**' in line:
            p = doc.add_paragraph()
            match = re.match(r'\*\*(.+?)\*\*\s*(.+)?', line)
            if match:
                run = p.add_run(match.group(1))
                run.bold = True
                if match.group(2):
                    p.add_run(' ' + match.group(2))
            i += 1
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Checkbox items
        if line.startswith('- [ ]'):
            text = line[6:]
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run('☐ ')
            add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        else:
            doc.add_paragraph()

        i += 1

    # Handle remaining table
    if in_table and table_rows:
        render_table(doc, table_rows)

    # Save
    output_path = os.path.join(script_dir, "Mining_DD_CoT_Framework.docx")
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")
    return output_path


def render_table(doc, rows):
    """Render a table in the Word document."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text

                # Bold and shade header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    # Add shading to header
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w'))
                    )
                    cell._tc.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph()


def add_formatted_text(paragraph, text):
    """Add text with inline formatting to a paragraph."""
    parts = []
    current_pos = 0

    # Find formatting markers
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))'

    for match in re.finditer(pattern, text):
        if match.start() > current_pos:
            parts.append(('normal', text[current_pos:match.start()]))

        matched_text = match.group()
        if matched_text.startswith('**') and matched_text.endswith('**'):
            parts.append(('bold', matched_text[2:-2]))
        elif matched_text.startswith('*') and matched_text.endswith('*'):
            parts.append(('italic', matched_text[1:-1]))
        elif matched_text.startswith('`') and matched_text.endswith('`'):
            parts.append(('code', matched_text[1:-1]))
        elif matched_text.startswith('['):
            # Handle links - extract text only
            link_match = re.match(r'\[([^\]]+)\]\([^)]+\)', matched_text)
            if link_match:
                parts.append(('normal', link_match.group(1)))

        current_pos = match.end()

    if current_pos < len(text):
        parts.append(('normal', text[current_pos:]))

    if not parts:
        paragraph.add_run(text)
        return

    for style, content in parts:
        run = paragraph.add_run(content)
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True
        elif style == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 0, 0)


if __name__ == "__main__":
    create_cot_framework_doc()
