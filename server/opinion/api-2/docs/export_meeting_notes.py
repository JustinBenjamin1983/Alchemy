"""
Export DD Tool Meeting Notes to Word Document
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def create_meeting_notes_doc():
    """Convert the markdown to a Word document."""

    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, "DD_Tool_Meeting_Notes.md")

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_rows = []

            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and all(set(c) <= set('-: ') for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            if table_rows:
                render_table(doc, table_rows)
            in_table = False
            table_rows = []
            continue

        # Main title
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Subtitle
        if line.startswith('## ') and 'Enhancement' in line:
            p = doc.add_paragraph(line[3:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(102, 102, 102)
            i += 1
            continue

        # Section headers (## 1. What...)
        if line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            i += 1
            continue

        # Subsection headers (### 3.1 ...)
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 76, 153)
            p.runs[0].font.size = Pt(12)
            i += 1
            continue

        # Horizontal rules
        if line.startswith('---'):
            p = doc.add_paragraph()
            run = p.add_run('─' * 70)
            run.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue

        # Bold labels like **What:** or **Why:**
        if line.startswith('**') and ':**' in line:
            p = doc.add_paragraph()
            match = re.match(r'\*\*(.+?)\*\*\s*(.+)?', line)
            if match:
                run = p.add_run(match.group(1))
                run.bold = True
                run.font.color.rgb = RGBColor(0, 76, 153)
                if match.group(2):
                    p.add_run(' ' + match.group(2))
            i += 1
            continue

        # Bullet points
        if line.startswith('- '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line)

        i += 1

    # Handle remaining table
    if in_table and table_rows:
        render_table(doc, table_rows)

    # Save
    output_path = os.path.join(script_dir, "DD_Tool_Meeting_Notes.docx")
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")
    return output_path


def render_table(doc, rows):
    """Render a table in the Word document."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w'))
                    )
                    cell._tc.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph()


def add_formatted_text(paragraph, text):
    """Add text with inline formatting to a paragraph."""
    parts = []
    current_pos = 0

    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'

    for match in re.finditer(pattern, text):
        if match.start() > current_pos:
            parts.append(('normal', text[current_pos:match.start()]))

        matched_text = match.group()
        if matched_text.startswith('**') and matched_text.endswith('**'):
            parts.append(('bold', matched_text[2:-2]))
        elif matched_text.startswith('*') and matched_text.endswith('*'):
            parts.append(('italic', matched_text[1:-1]))
        elif matched_text.startswith('`') and matched_text.endswith('`'):
            parts.append(('code', matched_text[1:-1]))

        current_pos = match.end()

    if current_pos < len(text):
        parts.append(('normal', text[current_pos:]))

    if not parts:
        paragraph.add_run(text)
        return

    for style, content in parts:
        run = paragraph.add_run(content)
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True
        elif style == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(10)


if __name__ == "__main__":
    create_meeting_notes_doc()
