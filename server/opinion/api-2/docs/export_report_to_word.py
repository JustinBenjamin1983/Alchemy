"""
Export DD Enhancement Summary Report to Word Document
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def create_word_report():
    """Convert the markdown report to a professional Word document."""

    # Read the markdown file
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, "DD_Enhancement_Summary_Report.md")

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)

    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)

    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)

    # Heading 3 style
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(51, 102, 153)

    # Process content line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines at start
        if i == 0 and not line:
            i += 1
            continue

        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # Start code block
                pass
            i += 1
            continue

        if in_code_block:
            # Add as code (monospace)
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue

        # Tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_rows = []

            # Parse table row
            cells = [c.strip() for c in line.split('|')[1:-1]]

            # Skip separator rows
            if cells and all(set(c) <= set('-: ') for c in cells):
                i += 1
                continue

            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table, render it
            if table_rows:
                render_table(doc, table_rows)
            in_table = False
            table_rows = []
            # Don't increment i, process this line
            continue

        # Main title
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_paragraph(line[2:], style='CustomTitle')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Subtitle
        if line.startswith('## ') and i < 5:
            p = doc.add_paragraph(line[3:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(102, 102, 102)
            i += 1
            continue

        # Section headers
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue

        if line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            i += 1
            continue

        # Horizontal rules
        if line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('─' * 80)
            p.runs[0].font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue

        # Bold metadata lines
        if line.startswith('**') and ':**' in line:
            p = doc.add_paragraph()
            # Parse bold label
            match = re.match(r'\*\*(.+?)\*\*\s*(.+)?', line)
            if match:
                run = p.add_run(match.group(1))
                run.bold = True
                if match.group(2):
                    p.add_run(' ' + match.group(2))
            i += 1
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            text = process_inline_formatting(line)
            p = doc.add_paragraph()
            add_formatted_text(p, text)
        else:
            # Empty line - add spacing
            doc.add_paragraph()

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        render_table(doc, table_rows)

    # Save document
    output_path = os.path.join(script_dir, "DD_Enhancement_Summary_Report.docx")
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")
    return output_path


def render_table(doc, rows):
    """Render a table in the Word document."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text

                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    # Add space after table
    doc.add_paragraph()


def process_inline_formatting(text):
    """Process inline markdown formatting markers."""
    # This returns the text with markers that add_formatted_text will process
    return text


def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, code) to a paragraph."""
    # Split on formatting markers
    # Handle **bold**, *italic*, `code`

    parts = []
    current_pos = 0

    # Find all formatting markers
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'

    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > current_pos:
            parts.append(('normal', text[current_pos:match.start()]))

        matched_text = match.group()
        if matched_text.startswith('**') and matched_text.endswith('**'):
            parts.append(('bold', matched_text[2:-2]))
        elif matched_text.startswith('*') and matched_text.endswith('*'):
            parts.append(('italic', matched_text[1:-1]))
        elif matched_text.startswith('`') and matched_text.endswith('`'):
            parts.append(('code', matched_text[1:-1]))

        current_pos = match.end()

    # Add remaining text
    if current_pos < len(text):
        parts.append(('normal', text[current_pos:]))

    # If no formatting found, just add the text
    if not parts:
        paragraph.add_run(text)
        return

    # Add formatted runs
    for style, content in parts:
        run = paragraph.add_run(content)
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True
        elif style == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(10)


if __name__ == "__main__":
    create_word_report()
